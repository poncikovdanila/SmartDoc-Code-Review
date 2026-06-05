"""Генератор шаблонов .docx с правильным оформлением.

Создаёт пустой документ с корректными настройками форматирования
в соответствии с переданными правилами (или дефолтными АГУ ФЦТиК).
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.checkers.docx_checker import _get_rules


def generate_template(custom_rules: dict[str, Any] | None = None) -> bytes:
    """Создаёт шаблон .docx по правилам и возвращает байты."""
    rules = _get_rules(custom_rules)
    r_font = rules["font_name"]
    r_size = rules["font_size_pt"]
    r_spacing = rules["line_spacing"]
    r_indent = rules["first_line_indent_cm"]
    r_margins = rules["margins_cm"]

    doc = Document()

    # ─── Настройка полей страницы ───
    for section in doc.sections:
        section.left_margin = Cm(r_margins["left"])
        section.right_margin = Cm(r_margins["right"])
        section.top_margin = Cm(r_margins["top"])
        section.bottom_margin = Cm(r_margins["bottom"])

    # ─── Настройка стиля Normal ───
    normal = doc.styles["Normal"]
    normal.font.name = r_font
    normal.font.size = Pt(r_size)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = r_spacing
    normal.paragraph_format.first_line_indent = Cm(r_indent)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Устанавливаем шрифт для всех Unicode-диапазонов
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), r_font)

    # ─── Настройка стилей заголовков ───
    for level in range(1, 4):
        style_name = f"Heading {level}"
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = r_font
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        if level == 1:
            style.font.size = Pt(r_size + 2)
        elif level == 2:
            style.font.size = Pt(r_size + 1)
        else:
            style.font.size = Pt(r_size)

    # ─── Титульный лист ───
    _add_centered(doc, "", r_font, r_size)  # пустая строка
    _add_centered(doc, "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ", r_font, r_size - 2)
    _add_centered(doc, "РОССИЙСКОЙ ФЕДЕРАЦИИ", r_font, r_size - 2)
    _add_centered(doc, "", r_font, r_size)
    _add_centered(doc, "АСТРАХАНСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ", r_font, r_size, bold=True)
    _add_centered(doc, 'имени В.Н. Татищева', r_font, r_size)
    _add_centered(doc, "", r_font, r_size)
    _add_centered(doc, "Факультет цифровых технологий и кибербезопасности", r_font, r_size)
    _add_centered(doc, "Кафедра информационных технологий и кибербезопасности", r_font, r_size)
    _add_centered(doc, "", r_font, r_size)
    _add_centered(doc, "", r_font, r_size)
    _add_centered(doc, "", r_font, r_size)
    _add_centered(doc, "ОТЧЁТ", r_font, r_size + 4, bold=True)
    _add_centered(doc, "по лабораторной работе №__", r_font, r_size)
    _add_centered(doc, "", r_font, r_size)
    _add_centered(doc, 'Тема: «___________________________»', r_font, r_size)
    _add_centered(doc, "", r_font, r_size)
    _add_centered(doc, 'по дисциплине «___________________________»', r_font, r_size)
    _add_centered(doc, "", r_font, r_size)
    _add_centered(doc, "", r_font, r_size)
    _add_centered(doc, "", r_font, r_size)

    # Правая часть: выполнил / проверил
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = r_spacing
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("Выполнил: студент группы ____\n____________________\n\nПроверил:\n____________________")
    run.font.name = r_font
    run.font.size = Pt(r_size)
    run.font.color.rgb = RGBColor(0, 0, 0)

    _add_centered(doc, "", r_font, r_size)
    _add_centered(doc, "", r_font, r_size)
    _add_centered(doc, "Астрахань — 2026", r_font, r_size)

    # Разрыв страницы
    doc.add_page_break()

    # ─── Содержание ───
    h = doc.add_heading("Содержание", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_heading(h, r_font, r_size + 2)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = r_spacing
    p.paragraph_format.first_line_indent = Cm(r_indent)
    run = p.add_run("[Здесь будет автоматическое оглавление — Ссылки → Оглавление в Word]")
    run.font.name = r_font
    run.font.size = Pt(r_size)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.italic = True

    doc.add_page_break()

    # ─── Введение ───
    h = doc.add_heading("Введение", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_heading(h, r_font, r_size + 2)
    _add_body(doc, "[Текст введения: актуальность, цель, задачи работы.]", r_font, r_size, r_spacing, r_indent)

    doc.add_page_break()

    # ─── Основная часть ───
    h = doc.add_heading("1 Теоретическая часть", level=1)
    _style_heading(h, r_font, r_size + 2)
    _add_body(doc, "[Теоретические сведения по теме работы.]", r_font, r_size, r_spacing, r_indent)

    h = doc.add_heading("1.1 Подраздел", level=2)
    _style_heading(h, r_font, r_size + 1)
    _add_body(doc, "[Текст подраздела.]", r_font, r_size, r_spacing, r_indent)

    doc.add_page_break()

    h = doc.add_heading("2 Практическая часть", level=1)
    _style_heading(h, r_font, r_size + 2)
    _add_body(doc, "[Описание практической части: ход работы, результаты.]", r_font, r_size, r_spacing, r_indent)

    # Пример подписи к рисунку
    _add_body(doc, "", r_font, r_size, r_spacing, r_indent)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = r_spacing
    run = p.add_run("[Здесь будет рисунок]")
    run.font.name = r_font
    run.font.size = Pt(r_size)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = r_spacing
    run = p.add_run("Рисунок 1 — Название рисунка")
    run.font.name = r_font
    run.font.size = Pt(r_size)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Пример подписи к таблице
    _add_body(doc, "", r_font, r_size, r_spacing, r_indent)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = r_spacing
    run = p.add_run("Таблица 1 — Название таблицы")
    run.font.name = r_font
    run.font.size = Pt(r_size)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Пример таблицы
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    headers = ["Параметр", "Значение", "Примечание"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.font.name = r_font
            run.font.size = Pt(r_size)
            run.bold = True

    doc.add_page_break()

    # ─── Заключение ───
    h = doc.add_heading("Заключение", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_heading(h, r_font, r_size + 2)
    _add_body(doc, "[Выводы по результатам работы.]", r_font, r_size, r_spacing, r_indent)

    doc.add_page_break()

    # ─── Список литературы ───
    h = doc.add_heading("Список литературы", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_heading(h, r_font, r_size + 2)

    for i in range(1, 4):
        _add_body(doc, f"{i}. [Автор И.О. Название источника. — Город: Издательство, Год. — С. 00–00.]",
                  r_font, r_size, r_spacing, r_indent)

    # ─── Сохраняем ───
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_centered(doc, text, font, size, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        run.font.name = font
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = bold


def _add_body(doc, text, font, size, spacing, indent):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.first_line_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _style_heading(heading, font, size):
    for run in heading.runs:
        run.font.name = font
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True

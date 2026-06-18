"""Дополнительные проверки нормоконтроля .docx: подписи, библиография, форматирование.

Что проверяется:
    1. Подписи к таблицам
    2. Подписи к рисункам
    3. Список литературы
    4. Выравнивание текста (по ширине)
    5. Лишние пустые строки
    6. Оформление заголовков
    7. Нумерация страниц
    8. Гиперссылки
    9. Цвет текста
   10. Оформление таблиц (шрифт, границы)
   11. Лишние пробелы
   12. Наличие оглавления
"""
from __future__ import annotations

import re
from typing import Any

from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Регулярки допускают секционную нумерацию (1.1, 2.3) и все виды тире.
TABLE_CAPTION_RE = re.compile(r"^Таблица\s+(\d+(?:\.\d+)*)\s*([—–\-])\s*(.+)", re.IGNORECASE)
FIGURE_CAPTION_RE = re.compile(r"^Рисунок\s+(\d+(?:\.\d+)*)\s*([—–\-])\s*(.+)", re.IGNORECASE)

# Заголовки раздела «Список литературы» — разные варианты, как пишут студенты.
BIBLIOGRAPHY_HEADERS = (
    "список литературы",
    "список использованных источников",
    "библиографический список",
)

# Источник в библиографии должен начинаться с «1.», «2.», ... либо «1)».
BIB_ENTRY_RE = re.compile(r"^\s*\d+\s*[.)]\s+\S")


def check_table_captions(document: DocxDocument) -> list[dict[str, Any]]:
    """Проверяет подписи к таблицам.

    Логика: проходим по элементам тела документа в порядке их появления
    (через xml-дерево), смотрим, что находится непосредственно перед каждой
    таблицей. Если это абзац — он должен соответствовать паттерну подписи.
    """
    issues: list[dict[str, Any]] = []
    body = document.element.body
    children = list(body.iterchildren())

    table_index = 0
    expected_number = 0

    for i, child in enumerate(children):
        if not child.tag.endswith("}tbl"):
            continue
        table_index += 1
        expected_number += 1

        # Ищем непустой абзац перед таблицей
        prev_paragraph_text = None
        for j in range(i - 1, -1, -1):
            prev = children[j]
            if prev.tag.endswith("}p"):
                text = "".join(prev.itertext()).strip()
                if text:
                    prev_paragraph_text = text
                    break
        if prev_paragraph_text is None:
            issues.append(
                {
                    "location": f"Таблица №{table_index} в документе",
                    "code": "TABLE_NO_CAPTION",
                    "message": "Перед таблицей отсутствует подпись",
                    "description": (
                        f"Перед таблицей должна быть подпись формата "
                        f"«Таблица {expected_number} — Название»"
                    ),
                    "severity": "high",
                    "expected": f"Таблица {expected_number} — Название",
                    "actual": "подпись отсутствует",
                }
            )
            continue

        match = TABLE_CAPTION_RE.match(prev_paragraph_text)
        if not match:
            issues.append(
                {
                    "location": f"Таблица №{table_index}: подпись",
                    "code": "TABLE_CAPTION_FORMAT",
                    "message": (
                        "Подпись таблицы не соответствует формату «Таблица N — Название»"
                    ),
                    "description": (
                        "Подпись таблицы оформляется как "
                        "«Таблица N — Название» (тире, не дефис)"
                    ),
                    "severity": "medium",
                    "expected": f"Таблица N — Название",
                    "actual": prev_paragraph_text[:80],
                }
            )
            continue

        dash = match.group(2)
        # Дефис — ошибка, тире (длинное или короткое) — допустимо
        if dash == "-":
            issues.append(
                {
                    "location": f"Таблица №{table_index}: подпись",
                    "code": "TABLE_DASH_WRONG",
                    "message": "В подписи используется дефис вместо тире",
                    "description": (
                        "Между номером и названием ставится тире «—» или «–», не дефис"
                    ),
                    "severity": "low",
                    "expected": "— или –",
                    "actual": dash,
                }
            )

    return issues


def check_figure_captions(document: DocxDocument) -> list[dict[str, Any]]:
    """Проверяет подписи к рисункам (искомый паттерн в тексте абзацев).

    К сожалению, python-docx не позволяет легко найти inline-изображения,
    привязанные к конкретному месту, поэтому идём от обратного: ищем абзацы,
    которые ВЫГЛЯДЯТ как подписи к рисункам, и проверяем их формат.
    """
    issues: list[dict[str, Any]] = []
    figure_caption_count = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        # Ищем варианты «Рис.», «Рисунок», «Figure»
        if not re.match(r"^Рис(\.|унок)\s+\d+", text, re.IGNORECASE):
            continue
        figure_caption_count += 1
        match = FIGURE_CAPTION_RE.match(text)
        if not match:
            issues.append(
                {
                    "location": f"Рисунок №{figure_caption_count}: подпись",
                    "code": "FIGURE_CAPTION_FORMAT",
                    "message": (
                        "Подпись рисунка не соответствует формату «Рисунок N — Название»"
                    ),
                    "description": (
                        "Подпись рисунка оформляется как "
                        "«Рисунок N — Название» (тире, не дефис; "
                        "слово «Рисунок» полностью, не «Рис.»)"
                    ),
                    "severity": "medium",
                    "expected": "Рисунок N — Название",
                    "actual": text[:80],
                }
            )
            continue
        dash = match.group(2)
        # Дефис — ошибка, тире (длинное или короткое) — допустимо
        if dash == "-":
            issues.append(
                {
                    "location": f"Рисунок №{figure_caption_count}: подпись",
                    "code": "FIGURE_DASH_WRONG",
                    "message": "В подписи рисунка используется дефис вместо тире",
                    "description": (
                        "Между номером и названием ставится тире «—» или «–», не дефис"
                    ),
                    "severity": "low",
                    "expected": "— или –",
                    "actual": dash,
                }
            )

    return issues


def check_bibliography(document: DocxDocument) -> list[dict[str, Any]]:
    """Проверяет наличие и базовое оформление списка литературы.

    Алгоритм:
        1. Ищем абзац-заголовок «Список литературы» / «Список использованных
           источников» / «Библиографический список».
        2. Если не нашли — это уже замечание (для академической работы список
           литературы обязателен).
        3. Если нашли — собираем все непустые абзацы после него до конца
           документа и проверяем, что каждый начинается с «N.» или «N)».
    """
    issues: list[dict[str, Any]] = []
    paragraphs = document.paragraphs
    biblio_start = None

    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip().lower().rstrip(".:")
        if text in BIBLIOGRAPHY_HEADERS:
            biblio_start = i
            break

    if biblio_start is None:
        # Не штрафуем, но информируем — у студента может быть и не положено
        # (например, рабочая записка). Уровень — low.
        issues.append(
            {
                "location": "Документ в целом",
                "code": "BIBLIOGRAPHY_MISSING",
                "message": "Не найден раздел «Список литературы»",
                "description": (
                    "В академической работе ожидается раздел «Список литературы» "
                    "или «Список использованных источников»"
                ),
                "severity": "low",
                "expected": "наличие раздела",
                "actual": "не найден",
            }
        )
        return issues

    # Проверяем, что после заголовка идут пронумерованные источники
    entries_found = 0
    incorrectly_formatted: list[tuple[int, str]] = []

    for paragraph in paragraphs[biblio_start + 1:]:
        text = paragraph.text.strip()
        if not text:
            continue
        # Если попался ещё один заголовок раздела — выходим
        if (
            text.lower().rstrip(".:") in BIBLIOGRAPHY_HEADERS
            or "приложение" in text.lower()[:15]
        ):
            break
        if BIB_ENTRY_RE.match(text):
            entries_found += 1
        else:
            # Это либо абзац-комментарий, либо просто кривая запись
            incorrectly_formatted.append((entries_found + 1, text[:80]))

    if entries_found == 0:
        issues.append(
            {
                "location": "Раздел «Список литературы»",
                "code": "BIBLIOGRAPHY_EMPTY",
                "message": "Раздел найден, но в нём нет пронумерованных источников",
                "description": (
                    "Каждый источник в списке литературы оформляется отдельным "
                    "пунктом и начинается с номера: «1. Иванов И. И. ...»"
                ),
                "severity": "high",
                "expected": "нумерованный список источников",
                "actual": "источники не найдены",
            }
        )
    elif incorrectly_formatted:
        # Сообщаем максимум о трёх таких — иначе отчёт раздуется
        for next_number, snippet in incorrectly_formatted[:3]:
            issues.append(
                {
                    "location": f"Список литературы, после источника №{next_number - 1}",
                    "code": "BIBLIOGRAPHY_ENTRY_FORMAT",
                    "message": "Запись не начинается с номера и точки",
                    "description": (
                        "Каждый источник оформляется как «N. Автор. Название…»"
                    ),
                    "severity": "medium",
                    "expected": "N. ...",
                    "actual": snippet,
                }
            )

    return issues


# ═══════ 4. Выравнивание текста по ширине ═══════

def _is_heading(paragraph: Paragraph) -> bool:
    name = (paragraph.style.name or "").lower()
    if "heading" in name or "заголов" in name or "toc" in name:
        return True
    if name.startswith("+") or "раздел" in name:
        return True
    return False


def check_text_alignment(document: DocxDocument) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    count = 0

    # Находим начало основного текста (после титульного листа)
    body_keywords = {"содержание", "введение", "оглавление"}
    body_start = 0
    for i, para in enumerate(document.paragraphs):
        text = para.text.strip().lower()
        style_name = (para.style.name or "").lower()
        if "heading" in style_name or text in body_keywords:
            body_start = i
            break

    for i, paragraph in enumerate(document.paragraphs):
        if i < body_start:
            continue
        text = paragraph.text.strip()
        if not text or len(text) < 20:
            continue
        if _is_heading(paragraph):
            continue
        style_name = (paragraph.style.name or "").lower()
        if style_name == "title":
            continue
        if TABLE_CAPTION_RE.match(text) or re.match(r"^Рис", text, re.IGNORECASE):
            continue
        alignment = paragraph.alignment
        if alignment is not None and alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            count += 1
            if count <= 5:
                preview = text[:50] + "…" if len(text) > 50 else text
                align_names = {
                    WD_ALIGN_PARAGRAPH.LEFT: "по левому краю",
                    WD_ALIGN_PARAGRAPH.CENTER: "по центру",
                    WD_ALIGN_PARAGRAPH.RIGHT: "по правому краю",
                }
                issues.append({
                    "location": f"Абзац: «{preview}»",
                    "code": "ALIGN_NOT_JUSTIFY",
                    "message": "Текст не выровнен по ширине",
                    "description": "Основной текст работы выравнивается по ширине (ГОСТ 7.32)",
                    "severity": "medium",
                    "expected": "по ширине (justify)",
                    "actual": align_names.get(alignment, str(alignment)),
                })
    if count > 5:
        issues.append({
            "location": "Документ в целом",
            "code": "ALIGN_NOT_JUSTIFY",
            "message": f"Ещё {count - 5} абзацев не выровнены по ширине",
            "description": "Основной текст выравнивается по ширине",
            "severity": "medium",
            "expected": "по ширине", "actual": f"{count} абзацев",
        })
    return issues


# ═══════ 5. Лишние пустые строки ═══════

def check_extra_blank_lines(document: DocxDocument) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    consecutive_empty = 0
    reported = 0
    for paragraph in document.paragraphs:
        if not paragraph.text.strip():
            consecutive_empty += 1
        else:
            if consecutive_empty >= 2 and reported < 5:
                preview = paragraph.text.strip()[:40]
                issues.append({
                    "location": f"Перед абзацем: «{preview}»",
                    "code": "EXTRA_BLANK_LINES",
                    "message": f"{consecutive_empty} пустых строк подряд",
                    "description": "Лишние пустые строки. Допускается не более одной.",
                    "severity": "low",
                    "expected": "не более 1", "actual": f"{consecutive_empty}",
                })
                reported += 1
            consecutive_empty = 0
    return issues


# ═══════ 6. Оформление заголовков ═══════

def check_headings(document: DocxDocument) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    for paragraph in document.paragraphs:
        if not _is_heading(paragraph):
            continue
        text = paragraph.text.strip()
        if not text:
            continue
        if text.endswith("."):
            issues.append({
                "location": f"Заголовок: «{text[:60]}»",
                "code": "HEADING_ENDS_WITH_DOT",
                "message": "Заголовок заканчивается точкой",
                "description": "По ГОСТ заголовки не заканчиваются точкой",
                "severity": "medium",
                "expected": "без точки", "actual": text[-10:],
            })
        # Проверяем жирность: стиль + ран-уровень
        # Если стиль жирный и раны НЕ переопределяют — ОК
        # Если ран явно выставляет bold=False — это ошибка
        style_bold = False
        s = paragraph.style
        while s:
            if s.font.bold:
                style_bold = True
                break
            s = s.base_style
        has_explicit_non_bold = any(
            run.text.strip() and run.bold is False
            for run in paragraph.runs
        )
        has_non_bold_runs = any(
            run.text.strip() and not run.bold
            for run in paragraph.runs
        )
        is_bold = style_bold and not has_explicit_non_bold or (
            not style_bold and not has_non_bold_runs
        )
        if not is_bold and paragraph.runs:
                issues.append({
                    "location": f"Заголовок: «{text[:60]}»",
                    "code": "HEADING_NOT_BOLD",
                    "message": "Заголовок не выделен жирным",
                    "description": "Заголовки оформляются полужирным начертанием",
                    "severity": "medium",
                    "expected": "полужирный", "actual": "обычный",
                })
    return issues


# ═══════ 7. Нумерация страниц ═══════

def check_page_numbers(document: DocxDocument) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    has_page_numbers = False
    for section in document.sections:
        for part in (section.footer, section.header):
            if part is None:
                continue
            xml = part._element.xml
            if "fldChar" in xml or "PAGE" in xml or "w:fldSimple" in xml:
                has_page_numbers = True
                break
        if has_page_numbers:
            break
    if not has_page_numbers:
        issues.append({
            "location": "Документ в целом",
            "code": "NO_PAGE_NUMBERS",
            "message": "Не найдена нумерация страниц",
            "description": "Страницы нумеруются арабскими цифрами (ГОСТ 7.32)",
            "severity": "medium",
            "expected": "нумерация страниц", "actual": "не найдена",
        })
    return issues


# ═══════ 8. Гиперссылки ═══════

def check_hyperlinks(document: DocxDocument) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    link_count = sum(1 for p in document.paragraphs if "w:hyperlink" in p._element.xml)
    if link_count > 0:
        issues.append({
            "location": "Документ в целом",
            "code": "HYPERLINKS_FOUND",
            "message": f"Найдено гиперссылок: {link_count}",
            "description": "Гиперссылки следует убрать или оформить обычным текстом",
            "severity": "low",
            "expected": "обычный текст", "actual": f"{link_count} гиперссылок",
        })
    return issues


# ═══════ 9. Цвет текста ═══════

def check_text_color(document: DocxDocument) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    colored = sum(
        1 for p in document.paragraphs for r in p.runs
        if r.text.strip() and r.font.color and r.font.color.rgb
        and r.font.color.rgb != RGBColor(0, 0, 0)
    )
    if colored > 0:
        issues.append({
            "location": "Документ в целом",
            "code": "COLORED_TEXT",
            "message": f"Найден цветной текст: {colored} фрагментов",
            "description": "Весь текст должен быть чёрного цвета",
            "severity": "medium",
            "expected": "чёрный", "actual": f"{colored} цветных фрагментов",
        })
    return issues


# ═══════ 10. Оформление таблиц ═══════

def check_table_formatting(document: DocxDocument) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    for idx, table in enumerate(document.tables, 1):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text.strip() and run.font.name and run.font.name != "Times New Roman":
                            issues.append({
                                "location": f"Таблица {idx}",
                                "code": "TABLE_FONT_MISMATCH",
                                "message": f"Шрифт «{run.font.name}», требуется Times New Roman",
                                "description": "Текст в таблицах — Times New Roman",
                                "severity": "medium",
                                "expected": "Times New Roman", "actual": run.font.name,
                            })
                            return issues
    return issues


# ═══════ 11. Лишние пробелы ═══════

def check_extra_spaces(document: DocxDocument) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    double_spaces = sum(len(re.findall(r"  +", p.text)) for p in document.paragraphs if p.text.strip())
    space_punct = sum(len(re.findall(r" +[,;.!?]", p.text)) for p in document.paragraphs if p.text.strip())
    if double_spaces > 0:
        issues.append({
            "location": "Документ в целом",
            "code": "DOUBLE_SPACES",
            "message": f"Двойных пробелов: {double_spaces}",
            "description": "В тексте не должно быть двойных пробелов",
            "severity": "low",
            "expected": "одинарные пробелы", "actual": f"{double_spaces}",
        })
    if space_punct > 0:
        issues.append({
            "location": "Документ в целом",
            "code": "SPACE_BEFORE_PUNCT",
            "message": f"Пробелов перед знаками препинания: {space_punct}",
            "description": "Перед запятой, точкой и т.д. пробел не ставится",
            "severity": "low",
            "expected": "без пробела", "actual": f"{space_punct}",
        })
    return issues


# ═══════ 12. Наличие оглавления ═══════

def check_table_of_contents(document: DocxDocument) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    has_toc = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip().lower()
        style_name = (paragraph.style.name or "").lower()
        if "toc" in style_name:
            has_toc = True
            break
        if text in ("содержание", "оглавление"):
            has_toc = True
            break
    body_xml = document.element.body.xml
    if "TOC" in body_xml:
        has_toc = True
    if not has_toc:
        issues.append({
            "location": "Документ в целом",
            "code": "NO_TABLE_OF_CONTENTS",
            "message": "Не найдено оглавление",
            "description": "В работе должно быть оглавление с номерами страниц",
            "severity": "low",
            "expected": "раздел «Содержание»", "actual": "не найдено",
        })
    return issues


# ═══════ 13. Интервалы до/после абзаца ═══════

def check_paragraph_spacing(document: DocxDocument) -> list[dict[str, Any]]:
    """Проверяет, что основной текст не имеет лишних интервалов Before/After."""
    issues: list[dict[str, Any]] = []
    count = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text or len(text) < 15:
            continue
        if _is_heading(paragraph):
            continue
        style_name = (paragraph.style.name or "").lower()
        if "title" in style_name or "toc" in style_name or "caption" in style_name:
            continue
        pf = paragraph.paragraph_format
        # Проверяем space_before и space_after
        before = pf.space_before
        after = pf.space_after
        # Допускаем None (наследование) и 0; флагим если > 6pt (0.5 строки)
        threshold = 76200  # 6pt в EMU
        if before is not None and before > threshold:
            count += 1
            if count <= 5:
                preview = text[:40] + "…" if len(text) > 40 else text
                issues.append({
                    "location": f"Абзац: «{preview}»",
                    "code": "PARA_SPACING_BEFORE",
                    "message": f"Интервал перед абзацем: {before / 12700:.0f} пт",
                    "description": "Основной текст не должен иметь дополнительного интервала перед абзацем",
                    "severity": "medium",
                    "expected": "0 пт", "actual": f"{before / 12700:.0f} пт",
                })
        if after is not None and after > threshold:
            count += 1
            if count <= 5:
                preview = text[:40] + "…" if len(text) > 40 else text
                issues.append({
                    "location": f"Абзац: «{preview}»",
                    "code": "PARA_SPACING_AFTER",
                    "message": f"Интервал после абзаца: {after / 12700:.0f} пт",
                    "description": "Основной текст не должен иметь дополнительного интервала после абзаца",
                    "severity": "medium",
                    "expected": "0 пт", "actual": f"{after / 12700:.0f} пт",
                })
    if count > 5:
        issues.append({
            "location": "Документ в целом",
            "code": "PARA_SPACING_BEFORE",
            "message": f"Ещё {count - 5} абзацев с лишними интервалами",
            "description": "Интервалы Before/After вместо межстрочного интервала",
            "severity": "medium",
            "expected": "0 пт", "actual": f"{count} абзацев",
        })
    return issues


# ═══════ 14. Нумерация разделов ═══════

SECTION_NUM_RE = re.compile(r"^(\d+(?:\.\d+)*)\s+")

def check_heading_numbering(document: DocxDocument) -> list[dict[str, Any]]:
    """Проверяет последовательность нумерации заголовков."""
    issues: list[dict[str, Any]] = []
    numbers_seen: list[str] = []

    for paragraph in document.paragraphs:
        if not _is_heading(paragraph):
            continue
        text = paragraph.text.strip()
        if not text:
            continue
        m = SECTION_NUM_RE.match(text)
        if not m:
            continue
        num = m.group(1)
        numbers_seen.append(num)

    # Проверяем последовательность
    prev_parts = []
    for num in numbers_seen:
        parts = [int(x) for x in num.split(".")]
        if prev_parts:
            # Проверяем, нет ли пропуска
            if len(parts) == 1:
                expected = prev_parts[0] + 1 if len(prev_parts) == 1 else prev_parts[0] + 1
                if parts[0] > expected:
                    issues.append({
                        "location": f"Заголовок «{num} ...»",
                        "code": "HEADING_NUM_GAP",
                        "message": f"Пропуск в нумерации разделов: после {'.'.join(str(x) for x in prev_parts)} идёт {num}",
                        "description": "Нумерация разделов должна быть последовательной",
                        "severity": "medium",
                        "expected": f"{expected}", "actual": num,
                    })
        prev_parts = parts

    return issues


# ═══════ 15. Иерархия заголовков ═══════

def check_heading_hierarchy(document: DocxDocument) -> list[dict[str, Any]]:
    """Проверяет, что нет перескоков уровней (например, Heading 1 → Heading 3)."""
    issues: list[dict[str, Any]] = []
    prev_level = 0

    for paragraph in document.paragraphs:
        style_name = paragraph.style.name or ""
        text = paragraph.text.strip()
        if not text:
            continue

        level = 0
        if style_name.startswith("Heading "):
            try:
                level = int(style_name.split()[-1])
            except ValueError:
                continue
        elif "раздел" in style_name.lower() and "под" not in style_name.lower():
            level = 1
        elif "подраздел" in style_name.lower():
            level = 2
        else:
            continue

        if prev_level > 0 and level > prev_level + 1:
            issues.append({
                "location": f"Заголовок: «{text[:50]}»",
                "code": "HEADING_LEVEL_SKIP",
                "message": f"Перескок уровня заголовка: с {prev_level} на {level}",
                "description": "Нельзя перескакивать уровни заголовков (например, с раздела сразу на пункт)",
                "severity": "medium",
                "expected": f"уровень {prev_level + 1}", "actual": f"уровень {level}",
            })
        prev_level = level

    return issues


# ═══════ 16. Ссылки на рисунки и таблицы в тексте ═══════

def check_cross_references(document: DocxDocument) -> list[dict[str, Any]]:
    """Проверяет, что каждый рисунок и таблица упомянуты в тексте."""
    issues: list[dict[str, Any]] = []

    # Собираем все подписи и обычный текст отдельно
    figures = []
    tables = []
    body_text_parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        fig_m = re.match(r"^Рисунок\s+(\d+(?:\.\d+)*)", text, re.IGNORECASE)
        if fig_m:
            figures.append(fig_m.group(1))
            continue
        tbl_m = re.match(r"^Таблица\s+(\d+(?:\.\d+)*)", text, re.IGNORECASE)
        if tbl_m:
            tables.append(tbl_m.group(1))
            continue
        # Только обычный текст (не подписи) для поиска ссылок
        body_text_parts.append(text)

    body_text = " ".join(body_text_parts).lower()

    # Проверяем ссылки на рисунки
    for num in figures:
        patterns = [
            f"рисунок {num}",
            f"рисунке {num}",
            f"рисунка {num}",
            f"рисунку {num}",
            f"рисунком {num}",
            f"рис. {num}",
            f"(рис. {num}",
        ]
        found = any(p in body_text for p in patterns)
        if not found:
            issues.append({
                "location": f"Рисунок {num}",
                "code": "FIGURE_NO_REFERENCE",
                "message": f"Рисунок {num} не упоминается в тексте",
                "description": "Каждый рисунок должен быть упомянут в тексте до его появления",
                "severity": "medium",
                "expected": f"ссылка на рисунок {num}", "actual": "не найдена",
            })

    # Проверяем ссылки на таблицы
    for num in tables:
        patterns = [
            f"таблица {num}",
            f"таблице {num}",
            f"таблицы {num}",
            f"таблицу {num}",
            f"таблицей {num}",
            f"табл. {num}",
            f"(табл. {num}",
        ]
        found = any(p in body_text for p in patterns)
        if not found:
            issues.append({
                "location": f"Таблица {num}",
                "code": "TABLE_NO_REFERENCE",
                "message": f"Таблица {num} не упоминается в тексте",
                "description": "Каждая таблица должна быть упомянута в тексте",
                "severity": "medium",
                "expected": f"ссылка на таблицу {num}", "actual": "не найдена",
            })

    return issues


# ═══════ 17. Формат приложений ═══════

def check_appendix_format(document: DocxDocument) -> list[dict[str, Any]]:
    """Проверяет формат заголовков приложений: ПРИЛОЖЕНИЕ А, ПРИЛОЖЕНИЕ Б, ..."""
    issues: list[dict[str, Any]] = []
    appendix_re = re.compile(r"^ПРИЛОЖЕНИЕ\s+([А-Я])", re.IGNORECASE)
    expected_letters = "АБВГДЕЖИКЛМНОПРСТУФХЦЧШЩЭЮЯ"
    found_letters = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower()
        # Пропускаем оглавление и обычные абзацы — только заголовки
        if "toc" in style_name:
            continue
        if not ("heading" in style_name or style_name.startswith("+")):
            continue
        m = appendix_re.match(text)
        if not m:
            continue
        letter = m.group(1).upper()
        found_letters.append(letter)

        # Проверяем, что приложение оформлено заглавными
        if not text.startswith("ПРИЛОЖЕНИЕ"):
            issues.append({
                "location": f"Приложение {letter}",
                "code": "APPENDIX_FORMAT",
                "message": "Слово «ПРИЛОЖЕНИЕ» должно быть заглавными буквами",
                "description": "Формат: «ПРИЛОЖЕНИЕ А — Название»",
                "severity": "medium",
                "expected": "ПРИЛОЖЕНИЕ", "actual": text[:20],
            })

    # Проверяем последовательность букв
    for i, letter in enumerate(found_letters):
        expected = expected_letters[i] if i < len(expected_letters) else "?"
        if letter != expected:
            issues.append({
                "location": f"Приложение {letter}",
                "code": "APPENDIX_ORDER",
                "message": f"Нарушена последовательность приложений: ожидалось {expected}, найдено {letter}",
                "description": "Приложения нумеруются буквами русского алфавита по порядку: А, Б, В, ...",
                "severity": "medium",
                "expected": expected, "actual": letter,
            })
            break

    return issues


# ═══════ 18. Формат библиографических записей ═══════

BIB_YEAR_RE = re.compile(r"\b(19|20)\d{2}\b")
BIB_PAGES_RE = re.compile(r"\d+\s*[сcС]\b\.?|\d+\s*p\b\.?", re.IGNORECASE)

def check_bibliography_format(document: DocxDocument) -> list[dict[str, Any]]:
    """Проверяет формат отдельных записей в списке литературы."""
    issues: list[dict[str, Any]] = []
    in_bib = False
    entry_count = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        text_lower = text.lower()
        if any(h in text_lower for h in BIBLIOGRAPHY_HEADERS):
            in_bib = True
            continue
        if in_bib:
            # Конец библиографии — новый раздел
            if _is_heading(paragraph) and not any(h in text_lower for h in BIBLIOGRAPHY_HEADERS):
                break
            m = BIB_ENTRY_RE.match(text)
            if not m:
                continue
            entry_count += 1
            # Проверяем наличие года
            if not BIB_YEAR_RE.search(text):
                issues.append({
                    "location": f"Источник №{entry_count}",
                    "code": "BIB_NO_YEAR",
                    "message": f"Не найден год издания в записи",
                    "description": "Каждый источник должен содержать год издания",
                    "severity": "low",
                    "expected": "год (19xx или 20xx)", "actual": text[:60],
                })

    return issues

"""Модуль автоисправления .docx по требованиям нормоконтроля АГУ (ФЦТиК).

Исправляет:
    * поля страницы → 3/1.5/2/2 см
    * шрифт → Times New Roman 14 пт
    * цвет текста → чёрный (во всех абзацах, заголовках, таблицах)
    * межстрочный интервал → 1.5
    * отступ первой строки → 1.25 см
    * выравнивание текста → по ширине (justify)
    * картинки → по центру
    * подписи к рисункам/таблицам → по центру, без отступа
    * пустые абзацы между картинкой и подписью → удаляются
    * заголовки → жирный, убрать точку в конце
    * лишние пустые строки → удаляются (2+ подряд → 1)
    * гиперссылки → чёрный цвет, без подчёркивания
    * шрифт в таблицах → Times New Roman
    * двойные пробелы → одинарные
    * пробелы перед знаками препинания → удаляются
"""
from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.checkers.docx_checker import _is_text_paragraph, _get_rules

BLACK = RGBColor(0, 0, 0)

# Паттерны подписей к рисункам и таблицам
CAPTION_RE = re.compile(
    r"^\s*(Рис(\.|унок)|Таблица|Figure|Table)\s*\d+",
    re.IGNORECASE,
)


def _set_run_font(run, font_name: str = "Times New Roman",
                  font_size_pt: float = 14.0):
    """Выставляет шрифт, размер и чёрный цвет на одном run."""
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.color.rgb = BLACK

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font_name)


def _paragraph_has_image(paragraph) -> bool:
    """Проверяет, содержит ли абзац встроенное изображение."""
    xml = paragraph._element.xml
    return "blipFill" in xml or "w:drawing" in xml


def _is_caption(paragraph) -> bool:
    """Проверяет, является ли абзац подписью к рисунку или таблице."""
    text = paragraph.text.strip()
    return bool(text and CAPTION_RE.match(text))


def _is_empty_paragraph(paragraph) -> bool:
    """Пустой абзац без текста и без картинок."""
    return not paragraph.text.strip() and not _paragraph_has_image(paragraph)


def autofix_docx(file_path: Path, custom_rules: dict | None = None) -> bytes:
    """Открывает .docx, исправляет нарушения, возвращает байты."""
    rules = _get_rules(custom_rules)
    r_font = rules["font_name"]
    r_size = rules["font_size_pt"]
    r_spacing = rules["line_spacing"]
    r_indent = Cm(rules["first_line_indent_cm"])
    r_margins = {k: Cm(v) for k, v in rules["margins_cm"].items()}

    document = Document(file_path)

    # 1. Поля страницы
    for section in document.sections:
        section.left_margin = r_margins["left"]
        section.right_margin = r_margins["right"]
        section.top_margin = r_margins["top"]
        section.bottom_margin = r_margins["bottom"]

    # 2. Стиль Normal
    try:
        normal = document.styles["Normal"]
        normal.font.name = r_font
        normal.font.size = Pt(r_size)
        normal.font.color.rgb = BLACK
    except KeyError:
        pass

    # 3. Заголовки — чёрный цвет
    for style_name in ("Heading 1", "Heading 2", "Heading 3",
                       "Heading 4", "Heading 5", "Heading 6"):
        try:
            style = document.styles[style_name]
            style.font.color.rgb = BLACK
            style.font.name = r_font
        except KeyError:
            pass

    # 4. Удаляем пустые абзацы между картинкой и подписью.
    #    Работаем на уровне XML — удаляем <w:p> из body.
    body = document.element.body
    paragraphs = document.paragraphs

    # Собираем индексы пустых абзацев, которые стоят между картинкой и подписью
    to_remove = []
    for i in range(len(paragraphs)):
        if not _paragraph_has_image(paragraphs[i]):
            continue
        # Нашли картинку. Смотрим вперёд: пустые абзацы → подпись
        j = i + 1
        empty_between = []
        while j < len(paragraphs) and _is_empty_paragraph(paragraphs[j]):
            empty_between.append(j)
            j += 1
        # Если после пустых идёт подпись — удаляем пустые
        if j < len(paragraphs) and _is_caption(paragraphs[j]) and empty_between:
            to_remove.extend(empty_between)

    # Удаляем в обратном порядке, чтобы индексы не сбились
    for idx in reversed(to_remove):
        elem = paragraphs[idx]._element
        elem.getparent().remove(elem)

    # 5. Проходим по всем абзацам (после удаления — обновляем список)
    for paragraph in document.paragraphs:
        # Картинки — по центру
        if _paragraph_has_image(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                _set_run_font(run, r_font, r_size)
            continue

        # Подписи к рисункам/таблицам — по центру, без отступа, минимальный
        # интервал сверху чтобы не отрывалась от картинки
        if _is_caption(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.line_spacing = r_spacing
            for run in paragraph.runs:
                _set_run_font(run, r_font, r_size)
            continue

        # Заголовки — чёрный цвет, не трогаем интервал
        style_name = (paragraph.style.name or "").lower()
        if "heading" in style_name or "заголов" in style_name:
            for run in paragraph.runs:
                _set_run_font(run, r_font, r_size)
            continue

        # Пропускаем пустые, но ставим чёрный цвет если есть run-ы
        if not _is_text_paragraph(paragraph):
            for run in paragraph.runs:
                run.font.color.rgb = BLACK
            continue

        # Основной текст
        pf = paragraph.paragraph_format
        pf.line_spacing = r_spacing
        pf.first_line_indent = r_indent
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for run in paragraph.runs:
            _set_run_font(run, r_font, r_size)

    # 6. Заголовки: жирный, убрать точку в конце
    for paragraph in document.paragraphs:
        style_name = (paragraph.style.name or "").lower()
        if "heading" not in style_name and "заголов" not in style_name:
            continue
        text = paragraph.text.strip()
        if text.endswith("."):
            # Убираем точку из последнего непустого run
            for run in reversed(paragraph.runs):
                if run.text.rstrip().endswith("."):
                    run.text = run.text.rstrip()[:-1]
                    break
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = BLACK

    # 7. Удаляем лишние пустые строки (2+ подряд → оставляем 1)
    paragraphs_list = document.paragraphs
    to_remove_blanks = []
    consecutive = 0
    for i, p in enumerate(paragraphs_list):
        if not p.text.strip() and not _paragraph_has_image(p):
            consecutive += 1
            if consecutive >= 2:
                to_remove_blanks.append(i)
        else:
            consecutive = 0
    for idx in reversed(to_remove_blanks):
        elem = paragraphs_list[idx]._element
        elem.getparent().remove(elem)

    # 8. Убираем форматирование гиперссылок (синий → чёрный, без подчёркивания)
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.font.color and run.font.color.rgb and run.font.color.rgb != BLACK:
                run.font.color.rgb = BLACK
            if run.font.underline:
                run.font.underline = False

    # 9. Исправляем шрифт в таблицах
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = r_font
                        if run.font.size and run.font.size.pt > 14:
                            run.font.size = Pt(r_size)
                        run.font.color.rgb = BLACK

    # 10. Убираем двойные пробелы и пробелы перед знаками препинания
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            import re as _re
            new_text = _re.sub(r"  +", " ", text)
            new_text = _re.sub(r" +([,;.!?])", r"\1", new_text)
            if new_text != text:
                run.text = new_text

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()

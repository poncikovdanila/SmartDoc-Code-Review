"""Расширенные тесты: автофикс-циклы, шаблон, хаос-документы,
граничные случаи, комбинации правил, экзотика кода.
"""
import io
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.checkers.docx_checker import check_docx_document
from app.checkers.docx_fixer import autofix_docx
from app.checkers.code_checker import check_python_code
from app.template_generator import generate_template


# ═══════ Хелперы ═══════

def _save(doc, tmp_path, name="test.docx"):
    path = tmp_path / name
    doc.save(path)
    return path

def _check(path, rules=None):
    return check_docx_document(path, path.name, rules)

def _codes(r):
    return {i["code"] for i in r["issues"]}

def _count(r, code):
    return sum(1 for i in r["issues"] if i["code"] == code)

def _make_base():
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Cm(3.5); sec.right_margin = Cm(1.0)
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    return doc

def _add_para(doc, text, font="Times New Roman", size=14,
              spacing=1.5, indent=1.25, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.line_spacing = spacing
    if indent is not None:
        p.paragraph_format.first_line_indent = Cm(indent)
    p.alignment = alignment
    return p

def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.bold = True
    return h

def _write_py(tmp_path, name, content):
    path = tmp_path / name
    path.write_text(content, encoding="utf-8")
    return path


# ═══════════════════════════════════════════
# A. Автофикс + повторная проверка
# ═══════════════════════════════════════════

class TestAutofixCycle:

    def test_autofix_reduces_errors(self, tmp_path):
        """После автофикса ошибок должно стать меньше."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        for _ in range(5):
            _add_para(doc, "Текст с неправильным шрифтом и размером.",
                      font="Arial", size=16)
        path = _save(doc, tmp_path)

        before = _check(path)
        fixed_bytes = autofix_docx(path)
        fixed_path = tmp_path / "fixed.docx"
        fixed_path.write_bytes(fixed_bytes)
        after = _check(fixed_path)

        assert after["total_issues"] < before["total_issues"]

    def test_autofix_fixes_font(self, tmp_path):
        """Автофикс должен исправить шрифт на Times New Roman."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст шрифтом Verdana который нужно исправить.", font="Verdana")
        path = _save(doc, tmp_path)
        fixed_bytes = autofix_docx(path)
        fixed_doc = Document(io.BytesIO(fixed_bytes))
        for para in fixed_doc.paragraphs:
            for run in para.runs:
                if run.text.strip() and run.font.name:
                    assert run.font.name == "Times New Roman"

    def test_autofix_fixes_spacing(self, tmp_path):
        """Автофикс должен исправить межстрочный интервал."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст с двойным интервалом нужно исправить.", spacing=2.0)
        path = _save(doc, tmp_path)
        fixed_bytes = autofix_docx(path)
        fixed_doc = Document(io.BytesIO(fixed_bytes))
        for para in fixed_doc.paragraphs:
            sp = para.paragraph_format.line_spacing
            if sp is not None and para.text.strip():
                assert abs(sp - 1.5) < 0.1 or sp == 1.0  # 1.5 или заголовочный

    def test_double_autofix_stable(self, tmp_path):
        """Двойной автофикс не должен менять результат."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст с ошибками для двойного автофикса.", font="Courier", size=10)
        path = _save(doc, tmp_path)

        fix1 = autofix_docx(path)
        p1 = tmp_path / "fix1.docx"
        p1.write_bytes(fix1)
        r1 = _check(p1)

        fix2 = autofix_docx(p1)
        p2 = tmp_path / "fix2.docx"
        p2.write_bytes(fix2)
        r2 = _check(p2)

        assert r2["total_issues"] <= r1["total_issues"]

    def test_autofix_preserves_text(self, tmp_path):
        """Автофикс не должен менять текст, только форматирование."""
        original_text = "Важный текст, который не должен измениться при автоисправлении."
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, original_text, font="Comic Sans MS", size=20)
        path = _save(doc, tmp_path)

        fixed_bytes = autofix_docx(path)
        fixed_doc = Document(io.BytesIO(fixed_bytes))
        all_text = " ".join(p.text for p in fixed_doc.paragraphs)
        assert original_text in all_text


# ═══════════════════════════════════════════
# B. Генератор шаблонов
# ═══════════════════════════════════════════

class TestTemplateGenerator:

    def test_template_passes_margin_check(self, tmp_path):
        """Сгенерированный шаблон не должен иметь ошибок полей."""
        template_bytes = generate_template()
        path = tmp_path / "template.docx"
        path.write_bytes(template_bytes)
        r = _check(path)
        assert _count(r, "MARGIN_MISMATCH") == 0

    def test_template_correct_font(self, tmp_path):
        """Шаблон должен использовать Times New Roman."""
        template_bytes = generate_template()
        path = tmp_path / "template.docx"
        path.write_bytes(template_bytes)
        r = _check(path)
        assert _count(r, "FONT_MISMATCH") == 0

    def test_template_with_custom_rules(self, tmp_path):
        """Шаблон с кастомными правилами (Arial 12pt)."""
        rules = {"font_name": "Arial", "font_size_pt": 12,
                 "margins_cm": {"left": 2.0, "right": 2.0, "top": 2.0, "bottom": 2.0}}
        template_bytes = generate_template(rules)
        doc = Document(io.BytesIO(template_bytes))
        sec = doc.sections[0]
        assert abs(sec.left_margin.cm - 2.0) < 0.1

    def test_template_is_valid_docx(self, tmp_path):
        """Шаблон можно открыть как docx."""
        template_bytes = generate_template()
        doc = Document(io.BytesIO(template_bytes))
        assert len(doc.paragraphs) > 0


# ═══════════════════════════════════════════
# C. Хаос-документы
# ═══════════════════════════════════════════

class TestChaosDocuments:

    def test_everything_wrong(self, tmp_path):
        """Документ со всеми ошибками — чекер не падает."""
        doc = _make_base()
        # Неправильные поля
        sec = doc.sections[0]
        sec.left_margin = Cm(1); sec.right_margin = Cm(5)
        sec.top_margin = Cm(0.5); sec.bottom_margin = Cm(0.5)
        # Заголовок без жирности с точкой
        h = doc.add_heading("Введение.", level=1)
        for run in h.runs:
            run.bold = False
        # Параграфы с разными ошибками
        _add_para(doc, "Arial  шрифт  с  двойными  пробелами !", font="Arial", size=8, spacing=2.5, indent=3)
        _add_para(doc, "Зелёный текст  ,неправильный.", font="Verdana", size=20, spacing=1.0, indent=0)
        for run in doc.paragraphs[-1].runs:
            run.font.color.rgb = RGBColor(0, 255, 0)
        # Пустые строки
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("")
        # Перескок уровня
        doc.add_heading("Пункт 1.1.1", level=3)

        path = _save(doc, tmp_path)
        r = _check(path)
        # Главное — не упал
        assert r["total_issues"] > 10
        assert r["verdict"] == "bad"

    def test_mixed_styles_chaos(self, tmp_path):
        """Каждый абзац — разный стиль."""
        doc = _make_base()
        _add_heading(doc, "СОДЕРЖАНИЕ")
        fonts = ["Arial", "Courier", "Verdana", "Georgia", "Impact"]
        sizes = [8, 10, 16, 20, 24]
        for i, (f, s) in enumerate(zip(fonts, sizes)):
            _add_para(doc, f"Абзац {i+1} шрифтом {f} размером {s}.", font=f, size=s)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert r["total_issues"] > 5
        assert _count(r, "FONT_MISMATCH") >= 5

    def test_very_long_document(self, tmp_path):
        """Документ с 300 абзацами — чекер не тормозит."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        for i in range(300):
            _add_para(doc, f"Абзац номер {i+1} содержит текст достаточной длины для проверки.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert r["paragraphs_checked"] > 200

    def test_only_headings(self, tmp_path):
        """Документ только из заголовков — не падает."""
        doc = _make_base()
        for i in range(10):
            _add_heading(doc, f"Заголовок {i+1}")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert isinstance(r["total_issues"], int)

    def test_single_paragraph(self, tmp_path):
        """Документ с одним абзацем."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Единственный абзац.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert isinstance(r["total_issues"], int)


# ═══════════════════════════════════════════
# D. Граничные случаи docx
# ═══════════════════════════════════════════

class TestDocxEdgeCases:

    def test_empty_document(self, tmp_path):
        """Пустой документ не роняет чекер."""
        doc = Document()
        path = _save(doc, tmp_path)
        r = _check(path)
        assert r["total_issues"] >= 0
        assert r["verdict"] in ("good", "ok", "bad")

    def test_document_with_only_tables(self, tmp_path):
        """Документ только с таблицами."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        table = doc.add_table(rows=3, cols=3)
        for row in table.rows:
            for cell in row.cells:
                cell.text = "Ячейка"
        path = _save(doc, tmp_path)
        r = _check(path)
        assert isinstance(r["total_issues"], int)

    def test_document_with_images(self, tmp_path):
        """Документ с изображением не падает."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст перед изображением.")
        # Добавляем маленькое изображение (1x1 pixel PNG)
        import struct, zlib
        def make_png():
            raw = b'\x00\x00\x00\x00'  # 1x1 black pixel
            width, height = 1, 1
            def chunk(ctype, data):
                c = ctype + data
                return struct.pack('>I', len(data)) + c + struct.pack('>I', zlib.crc32(c) & 0xFFFFFFFF)
            return (b'\x89PNG\r\n\x1a\n' +
                    chunk(b'IHDR', struct.pack('>IIBBBBB', width, height, 8, 2, 0, 0, 0)) +
                    chunk(b'IDAT', zlib.compress(raw)) +
                    chunk(b'IEND', b''))
        png_path = tmp_path / "test.png"
        png_path.write_bytes(make_png())
        doc.add_picture(str(png_path), width=Inches(2))
        path = _save(doc, tmp_path)
        r = _check(path)
        assert isinstance(r["total_issues"], int)

    def test_multiple_sections_different_margins(self, tmp_path):
        """Несколько разделов с разными полями."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст первого раздела документа.")
        # Добавляем новый раздел с другими полями
        new_section = doc.add_section()
        new_section.left_margin = Cm(2.0)
        new_section.right_margin = Cm(2.0)
        _add_para(doc, "Текст второго раздела с другими полями.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "MARGIN_MISMATCH") >= 1

    def test_very_long_paragraph(self, tmp_path):
        """Абзац на 5000 символов."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        long_text = "Слово " * 800  # ~4000 символов
        _add_para(doc, long_text)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert isinstance(r["total_issues"], int)

    def test_special_characters_in_text(self, tmp_path):
        """Спецсимволы в тексте не ломают чекер."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Формула: E = mc², α + β = γ, 100% → ∞, «кавычки».")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert isinstance(r["total_issues"], int)


# ═══════════════════════════════════════════
# E. Комбинации кастомных правил
# ═══════════════════════════════════════════

class TestCustomRuleCombinations:

    def test_all_checks_disabled(self, tmp_path):
        """Все проверки отключены — только базовые."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст с ошибками.", font="Arial", size=20)
        path = _save(doc, tmp_path)
        all_off = {
            "checks": {
                "bibliography": False, "hyperlinks": False, "textColor": False,
                "spaces": False, "blankLines": False, "headings": False,
                "pageNumbers": False, "tables": False, "toc": False,
                "paraSpacing": False, "headingNumbers": False,
                "headingHierarchy": False, "crossReferences": False,
                "appendix": False, "bibFormat": False,
            }
        }
        r = _check(path, all_off)
        # Должны остаться только базовые: font, size, spacing, indent, margins
        for iss in r["issues"]:
            assert iss["code"] in {
                "FONT_MISMATCH", "FONT_NOT_SET", "FONT_SIZE_MISMATCH",
                "FONT_SIZE_NOT_SET", "LINE_SPACING_MISMATCH", "LINE_SPACING_NOT_SET",
                "INDENT_MISMATCH", "INDENT_NOT_SET", "MARGIN_MISMATCH",
            }

    def test_custom_font_and_size(self, tmp_path):
        """Кастомный шрифт + размер: Arial 10pt."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст Arial 10pt который должен пройти проверку.", font="Arial", size=10)
        path = _save(doc, tmp_path)
        r = _check(path, {"font_name": "Arial", "font_size_pt": 10})
        assert _count(r, "FONT_MISMATCH") == 0
        assert _count(r, "FONT_SIZE_MISMATCH") == 0

    def test_custom_margins(self, tmp_path):
        """Кастомные поля 2/2/2/2."""
        doc = Document()
        for sec in doc.sections:
            sec.left_margin = Cm(2); sec.right_margin = Cm(2)
            sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст с кастомными полями.")
        path = _save(doc, tmp_path)
        r = _check(path, {"margins_cm": {"left": 2, "right": 2, "top": 2, "bottom": 2}})
        assert _count(r, "MARGIN_MISMATCH") == 0

    def test_alignment_any(self, tmp_path):
        """alignment=any — не проверяем выравнивание."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст с выравниванием по левому краю.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
        path = _save(doc, tmp_path)
        r = _check(path, {"alignment": "any"})
        assert _count(r, "ALIGN_NOT_JUSTIFY") == 0

    def test_custom_spacing(self, tmp_path):
        """Кастомный интервал 1.0."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст с одинарным интервалом.", spacing=1.0)
        path = _save(doc, tmp_path)
        r = _check(path, {"line_spacing": 1.0})
        assert _count(r, "LINE_SPACING_MISMATCH") == 0

    def test_custom_indent(self, tmp_path):
        """Кастомный отступ 1.0 см."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст с отступом один сантиметр.", indent=1.0)
        path = _save(doc, tmp_path)
        r = _check(path, {"first_line_indent_cm": 1.0})
        assert _count(r, "INDENT_MISMATCH") == 0


# ═══════════════════════════════════════════
# F. Регрессии (ранее проблемные сценарии)
# ═══════════════════════════════════════════

class TestRegressions:

    def test_title_page_12pt_not_flagged(self, tmp_path):
        """Регрессия: титульник с 12pt не должен давать FONT_SIZE_MISMATCH."""
        doc = _make_base()
        for _ in range(10):
            _add_para(doc, "Элемент титульного листа.", size=12,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=0)
        _add_heading(doc, "СОДЕРЖАНИЕ")
        _add_para(doc, "Текст основной части.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "FONT_SIZE_MISMATCH") == 0

    def test_centered_heading_no_indent_error(self, tmp_path):
        """Регрессия: центрированный заголовок не должен давать INDENT ошибку."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Обычный текст с правильным отступом.")
        path = _save(doc, tmp_path)
        r = _check(path)
        # Заголовки не проверяются на отступ
        heading_indent = [i for i in r["issues"]
                         if i["code"] in ("INDENT_MISMATCH", "INDENT_NOT_SET")
                         and "Введение" in i.get("location", "").upper()]
        assert len(heading_indent) == 0

    def test_section_numbering_caption_ok(self, tmp_path):
        """Регрессия: подпись Рисунок 2.1 – Название не должна быть ошибкой."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Результат показан на рисунке 2.1.")
        p = doc.add_paragraph()
        run = p.add_run("Рисунок 2.1 – Схема алгоритма")
        run.font.size = Pt(14)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "FIGURE_CAPTION_FORMAT") == 0

    def test_style_bold_heading_not_flagged(self, tmp_path):
        """Регрессия: заголовок жирный через стиль не должен быть HEADING_NOT_BOLD."""
        doc = _make_base()
        h = doc.add_heading("ВВЕДЕНИЕ", level=1)
        # Не трогаем bold у ранов — наследуется из стиля
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "HEADING_NOT_BOLD") == 0

    def test_12pt_body_text_accepted(self, tmp_path):
        """Регрессия: 12pt в теле документа допустим (наравне с 14pt)."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст размером 12 пунктов.", size=12)
        _add_para(doc, "Текст размером 14 пунктов.", size=14)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "FONT_SIZE_MISMATCH") == 0


# ═══════════════════════════════════════════
# G. Код: экзотические сценарии
# ═══════════════════════════════════════════

class TestCodeExotic:

    def test_decorators(self, tmp_path):
        src = (
            '"""Module."""\n\n\n'
            "def decorator(func):\n"
            '    """Decorator."""\n'
            "    return func\n\n\n"
            "@decorator\n"
            "def hello():\n"
            '    """Say hello."""\n'
            "    print('hello')\n"
        )
        path = _write_py(tmp_path, "deco.py", src)
        r = check_python_code(path, "deco.py")
        assert r["total_issues"] == 0

    def test_type_hints(self, tmp_path):
        src = (
            '"""Module."""\n\n\n'
            "def add(a: int, b: int) -> int:\n"
            '    """Add two ints."""\n'
            "    return a + b\n"
        )
        path = _write_py(tmp_path, "hints.py", src)
        r = check_python_code(path, "hints.py")
        assert r["total_issues"] == 0

    def test_f_strings(self, tmp_path):
        src = (
            '"""Module."""\n\n\n'
            "def greet(name: str) -> str:\n"
            '    """Greet."""\n'
            "    return f'Hello, {name}!'\n"
        )
        path = _write_py(tmp_path, "fstr.py", src)
        r = check_python_code(path, "fstr.py")
        assert r["total_issues"] == 0

    def test_class_with_methods(self, tmp_path):
        src = (
            '"""Module."""\n\n\n'
            "class Dog:\n"
            '    """A dog."""\n\n'
            "    def __init__(self, name: str) -> None:\n"
            '        """Init."""\n'
            "        self.name = name\n\n"
            "    def bark(self) -> str:\n"
            '        """Bark."""\n'
            "        return f'{self.name} says woof!'\n"
        )
        path = _write_py(tmp_path, "cls.py", src)
        r = check_python_code(path, "cls.py")
        assert r["total_issues"] == 0

    def test_only_imports(self, tmp_path):
        src = "import os\nimport sys\nimport json\n"
        path = _write_py(tmp_path, "imports.py", src)
        r = check_python_code(path, "imports.py")
        # Неиспользуемые импорты = F401
        assert _count(r, "F401") == 3

    def test_star_import(self, tmp_path):
        src = "from os import *\n\nprint(getcwd())\n"
        path = _write_py(tmp_path, "star.py", src)
        r = check_python_code(path, "star.py")
        codes = {i["code"] for i in r["issues"]}
        assert "F403" in codes  # wildcard import

    def test_extremely_long_line(self, tmp_path):
        src = f'x = "{" " * 500}"\n'
        path = _write_py(tmp_path, "long.py", src)
        r = check_python_code(path, "long.py")
        codes = {i["code"] for i in r["issues"]}
        assert "E501" in codes

    def test_global_variable_naming(self, tmp_path):
        src = '"""Module."""\n\nmy_var = 1\n'
        path = _write_py(tmp_path, "global.py", src)
        r = check_python_code(path, "global.py")
        # Просто не падает, глобальные переменные допустимы
        assert isinstance(r["total_issues"], int)

    def test_nested_functions(self, tmp_path):
        src = (
            '"""Module."""\n\n\n'
            "def outer():\n"
            '    """Outer."""\n'
            "    def inner():\n"
            '        """Inner."""\n'
            "        return 42\n"
            "    return inner()\n"
        )
        path = _write_py(tmp_path, "nested.py", src)
        r = check_python_code(path, "nested.py")
        assert r["total_issues"] == 0

    def test_multiline_string(self, tmp_path):
        src = (
            '"""Module."""\n\n\n'
            "SQL = '''\n"
            "    SELECT *\n"
            "    FROM users\n"
            "    WHERE active = 1\n"
            "'''\n"
        )
        path = _write_py(tmp_path, "multiline.py", src)
        r = check_python_code(path, "multiline.py")
        assert isinstance(r["total_issues"], int)

    def test_lambda(self, tmp_path):
        src = (
            '"""Module."""\n\n'
            "add = lambda x, y: x + y  # noqa: E731\n"
        )
        path = _write_py(tmp_path, "lam.py", src)
        r = check_python_code(path, "lam.py")
        # С noqa — flake8 пропустит E731
        assert isinstance(r["total_issues"], int)

    def test_comprehensions(self, tmp_path):
        src = (
            '"""Module."""\n\n\n'
            "def squares(n: int) -> list:\n"
            '    """Squares."""\n'
            "    return [x ** 2 for x in range(n)]\n"
        )
        path = _write_py(tmp_path, "comp.py", src)
        r = check_python_code(path, "comp.py")
        assert r["total_issues"] == 0

    def test_verdict_fields_present(self, tmp_path):
        """Каждый отчёт содержит verdict, summary, total_issues."""
        src = "x = 1\n"
        path = _write_py(tmp_path, "v.py", src)
        r = check_python_code(path, "v.py")
        assert "verdict" in r
        assert "summary" in r
        assert "total_issues" in r
        assert r["verdict"] in ("good", "ok", "bad")

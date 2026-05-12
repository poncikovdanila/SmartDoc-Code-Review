"""40 тестов проверки различных .docx документов.

Каждый тест создаёт документ с определённым набором свойств
и проверяет, что docx_checker и docx_fixer реагируют корректно.
"""
import io
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.checkers.docx_checker import check_docx_document, _get_rules
from app.checkers.docx_fixer import autofix_docx


# ═══════ Хелперы ═══════

def _make_doc(tmp_path: Path, filename: str = "test.docx", **kwargs) -> Path:
    """Создаёт .docx с настраиваемыми параметрами.

    kwargs:
        font_name, font_size, line_spacing, indent, text,
        margins (dict: left/right/top/bottom),
        paragraphs (list[str]), add_bibliography (bool)
    """
    doc = Document()
    sec = doc.sections[0]

    margins = kwargs.get("margins", {"left": 3.0, "right": 1.5, "top": 2.0, "bottom": 2.0})
    sec.left_margin = Cm(margins.get("left", 3.0))
    sec.right_margin = Cm(margins.get("right", 1.5))
    sec.top_margin = Cm(margins.get("top", 2.0))
    sec.bottom_margin = Cm(margins.get("bottom", 2.0))

    normal = doc.styles["Normal"]
    normal.font.name = kwargs.get("font_name", "Times New Roman")
    normal.font.size = Pt(kwargs.get("font_size", 14))

    texts = kwargs.get("paragraphs", [kwargs.get("text", "Пример абзаца основного текста документа.")])
    for text in texts:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = kwargs.get("font_name", "Times New Roman")
        run.font.size = Pt(kwargs.get("font_size", 14))
        p.paragraph_format.line_spacing = kwargs.get("line_spacing", 1.5)
        p.paragraph_format.first_line_indent = Cm(kwargs.get("indent", 1.25))

    if kwargs.get("add_bibliography", True):
        bh = doc.add_paragraph("Список литературы")
        bh.paragraph_format.line_spacing = 1.5
        bh.paragraph_format.first_line_indent = Cm(1.25)
        bi = doc.add_paragraph("1. Иванов И. И. Основы информатики. — М.: Наука, 2020.")
        bi.paragraph_format.line_spacing = 1.5
        bi.paragraph_format.first_line_indent = Cm(1.25)

    path = tmp_path / filename
    doc.save(path)
    return path


def _check(path: Path, rules=None) -> dict:
    return check_docx_document(path, path.name, rules)


def _fix_and_recheck(path: Path, rules=None) -> dict:
    fixed_bytes = autofix_docx(path, rules)
    fixed_path = path.parent / "fixed.docx"
    fixed_path.write_bytes(fixed_bytes)
    return check_docx_document(fixed_path, "fixed.docx", rules)


# ═══════ Полностью корректные документы ═══════

class TestCompliantDocs:
    def test_standard_agu_doc(self, tmp_path):
        path = _make_doc(tmp_path)
        r = _check(path)
        assert r["total_issues"] == 0

    def test_multiple_paragraphs(self, tmp_path):
        path = _make_doc(tmp_path, paragraphs=[
            "Первый абзац текста работы.",
            "Второй абзац с продолжением.",
            "Третий абзац завершает мысль.",
        ])
        r = _check(path)
        assert r["total_issues"] == 0

    def test_long_paragraph(self, tmp_path):
        text = "Слово " * 200
        path = _make_doc(tmp_path, text=text.strip())
        r = _check(path)
        font_issues = [i for i in r["issues"] if "FONT" in i["code"]]
        assert len(font_issues) == 0

    def test_paragraphs_checked_count(self, tmp_path):
        path = _make_doc(tmp_path, paragraphs=["А", "Б", "В"])
        r = _check(path)
        # 3 основных + 2 библиография = 5
        assert r["paragraphs_checked"] == 5


# ═══════ Ошибки шрифта ═══════

class TestFontErrors:
    def test_arial_instead_of_tnr(self, tmp_path):
        path = _make_doc(tmp_path, font_name="Arial")
        r = _check(path)
        codes = {i["code"] for i in r["issues"]}
        assert "FONT_MISMATCH" in codes

    def test_courier_font(self, tmp_path):
        path = _make_doc(tmp_path, font_name="Courier New")
        r = _check(path)
        codes = {i["code"] for i in r["issues"]}
        assert "FONT_MISMATCH" in codes

    def test_font_size_12_instead_of_14(self, tmp_path):
        path = _make_doc(tmp_path, font_size=12)
        r = _check(path)
        codes = {i["code"] for i in r["issues"]}
        assert "FONT_SIZE_MISMATCH" in codes

    def test_font_size_16(self, tmp_path):
        path = _make_doc(tmp_path, font_size=16)
        r = _check(path)
        codes = {i["code"] for i in r["issues"]}
        assert "FONT_SIZE_MISMATCH" in codes

    def test_font_size_10(self, tmp_path):
        path = _make_doc(tmp_path, font_size=10)
        r = _check(path)
        size_issues = [i for i in r["issues"] if i["code"] == "FONT_SIZE_MISMATCH"]
        assert len(size_issues) > 0

    def test_font_mismatch_is_high_severity(self, tmp_path):
        path = _make_doc(tmp_path, font_name="Arial")
        r = _check(path)
        font_iss = [i for i in r["issues"] if i["code"] == "FONT_MISMATCH"]
        assert all(i["severity"] == "high" for i in font_iss)


# ═══════ Ошибки интервала ═══════

class TestSpacingErrors:
    def test_single_spacing(self, tmp_path):
        path = _make_doc(tmp_path, line_spacing=1.0)
        r = _check(path)
        codes = {i["code"] for i in r["issues"]}
        assert "LINE_SPACING_MISMATCH" in codes

    def test_double_spacing(self, tmp_path):
        path = _make_doc(tmp_path, line_spacing=2.0)
        r = _check(path)
        codes = {i["code"] for i in r["issues"]}
        assert "LINE_SPACING_MISMATCH" in codes

    def test_spacing_mismatch_is_high(self, tmp_path):
        path = _make_doc(tmp_path, line_spacing=1.0)
        r = _check(path)
        sp = [i for i in r["issues"] if i["code"] == "LINE_SPACING_MISMATCH"]
        assert all(i["severity"] == "high" for i in sp)


# ═══════ Ошибки отступа ═══════

class TestIndentErrors:
    def test_no_indent(self, tmp_path):
        path = _make_doc(tmp_path, indent=0)
        r = _check(path)
        codes = {i["code"] for i in r["issues"]}
        assert "INDENT_MISMATCH" in codes or "INDENT_NOT_SET" in codes

    def test_large_indent(self, tmp_path):
        path = _make_doc(tmp_path, indent=2.5)
        r = _check(path)
        codes = {i["code"] for i in r["issues"]}
        assert "INDENT_MISMATCH" in codes

    def test_small_indent(self, tmp_path):
        path = _make_doc(tmp_path, indent=0.5)
        r = _check(path)
        codes = {i["code"] for i in r["issues"]}
        assert "INDENT_MISMATCH" in codes


# ═══════ Ошибки полей ═══════

class TestMarginErrors:
    def test_narrow_left_margin(self, tmp_path):
        path = _make_doc(tmp_path, margins={"left": 2.0, "right": 1.5, "top": 2.0, "bottom": 2.0})
        r = _check(path)
        margin_iss = [i for i in r["issues"] if i["code"] == "MARGIN_MISMATCH"]
        assert len(margin_iss) > 0

    def test_wide_right_margin(self, tmp_path):
        path = _make_doc(tmp_path, margins={"left": 3.0, "right": 3.0, "top": 2.0, "bottom": 2.0})
        r = _check(path)
        margin_iss = [i for i in r["issues"] if i["code"] == "MARGIN_MISMATCH"]
        assert len(margin_iss) > 0

    def test_all_margins_wrong(self, tmp_path):
        path = _make_doc(tmp_path, margins={"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0})
        r = _check(path)
        margin_iss = [i for i in r["issues"] if i["code"] == "MARGIN_MISMATCH"]
        assert len(margin_iss) == 4

    def test_margin_expected_and_actual(self, tmp_path):
        path = _make_doc(tmp_path, margins={"left": 2.0, "right": 1.5, "top": 2.0, "bottom": 2.0})
        r = _check(path)
        margin_iss = [i for i in r["issues"] if i["code"] == "MARGIN_MISMATCH"]
        for iss in margin_iss:
            assert "expected" in iss
            assert "actual" in iss


# ═══════ Множественные ошибки ═══════

class TestMultipleErrors:
    def test_all_wrong(self, tmp_path):
        path = _make_doc(tmp_path, font_name="Arial", font_size=12,
                         line_spacing=1.0, indent=0.5,
                         margins={"left": 2.0, "right": 2.0, "top": 2.5, "bottom": 2.5},
                         add_bibliography=False)
        r = _check(path)
        codes = {i["code"] for i in r["issues"]}
        assert "FONT_MISMATCH" in codes
        assert "FONT_SIZE_MISMATCH" in codes
        assert "LINE_SPACING_MISMATCH" in codes
        assert "INDENT_MISMATCH" in codes
        assert "MARGIN_MISMATCH" in codes
        assert "BIBLIOGRAPHY_MISSING" in codes

    def test_only_font_wrong(self, tmp_path):
        path = _make_doc(tmp_path, font_name="Verdana")
        r = _check(path)
        non_font_codes = {i["code"] for i in r["issues"]} - {"FONT_MISMATCH"}
        assert "FONT_SIZE_MISMATCH" not in non_font_codes
        assert "MARGIN_MISMATCH" not in non_font_codes


# ═══════ Автоисправление ═══════

class TestDocxAutofix:
    def test_fixes_font(self, tmp_path):
        path = _make_doc(tmp_path, font_name="Arial")
        fixed = autofix_docx(path)
        doc = Document(io.BytesIO(fixed))
        for p in doc.paragraphs:
            for run in p.runs:
                if run.text.strip():
                    assert run.font.name == "Times New Roman"

    def test_fixes_margins(self, tmp_path):
        path = _make_doc(tmp_path, margins={"left": 2.0, "right": 2.0, "top": 2.5, "bottom": 2.5})
        fixed = autofix_docx(path)
        doc = Document(io.BytesIO(fixed))
        sec = doc.sections[0]
        assert abs(sec.left_margin.cm - 3.0) < 0.1
        assert abs(sec.right_margin.cm - 1.5) < 0.1

    def test_fixes_font_size(self, tmp_path):
        path = _make_doc(tmp_path, font_size=12)
        fixed = autofix_docx(path)
        doc = Document(io.BytesIO(fixed))
        for p in doc.paragraphs:
            for run in p.runs:
                if run.text.strip() and run.font.size:
                    assert abs(run.font.size.pt - 14.0) < 0.5

    def test_autofix_reduces_issues(self, tmp_path):
        path = _make_doc(tmp_path, font_name="Arial", font_size=12,
                         margins={"left": 2.0, "right": 2.0, "top": 2.5, "bottom": 2.5})
        before = _check(path)
        after = _fix_and_recheck(path)
        assert after["total_issues"] < before["total_issues"]

    def test_autofix_with_custom_rules(self, tmp_path):
        path = _make_doc(tmp_path, font_name="Verdana", font_size=10)
        custom = {"font_name": "Arial", "font_size_pt": 12}
        fixed = autofix_docx(path, custom)
        doc = Document(io.BytesIO(fixed))
        for p in doc.paragraphs:
            for run in p.runs:
                if run.text.strip():
                    assert run.font.name == "Arial"

    def test_autofix_returns_valid_docx(self, tmp_path):
        path = _make_doc(tmp_path, font_name="Arial")
        fixed = autofix_docx(path)
        assert isinstance(fixed, bytes)
        doc = Document(io.BytesIO(fixed))
        assert len(doc.paragraphs) > 0


# ═══════ Пользовательские правила ═══════

class TestCustomRulesIntegration:
    def test_arial_12pt_passes_with_matching_rules(self, tmp_path):
        path = _make_doc(tmp_path, font_name="Arial", font_size=12)
        r = _check(path, {"font_name": "Arial", "font_size_pt": 12})
        font_iss = [i for i in r["issues"] if "FONT" in i["code"]]
        assert len(font_iss) == 0

    def test_custom_margins(self, tmp_path):
        path = _make_doc(tmp_path, margins={"left": 2.5, "right": 1.0, "top": 2.0, "bottom": 2.0})
        r_default = _check(path)
        r_custom = _check(path, {"margins_cm": {"left": 2.5, "right": 1.0}})
        assert r_custom["total_issues"] < r_default["total_issues"]

    def test_preset_agu(self, tmp_path):
        path = _make_doc(tmp_path)
        r = _check(path, {"preset": "agu"})
        assert r["total_issues"] == 0


# ═══════ Структура отчёта ═══════

class TestDocxReportStructure:
    def test_report_has_required_fields(self, tmp_path):
        path = _make_doc(tmp_path)
        r = _check(path)
        assert r["file_type"] == "docx"
        assert "total_issues" in r
        assert "summary" in r
        assert "issues" in r
        assert "paragraphs_checked" in r

    def test_summary_matches_total(self, tmp_path):
        path = _make_doc(tmp_path, font_name="Arial", font_size=12)
        r = _check(path)
        s = r["summary"]
        assert s["high"] + s["medium"] + s["low"] == r["total_issues"]

    def test_issue_has_location(self, tmp_path):
        path = _make_doc(tmp_path, font_name="Arial")
        r = _check(path)
        for iss in r["issues"]:
            if iss["code"] not in ("MARGIN_MISMATCH", "BIBLIOGRAPHY_MISSING"):
                assert "location" in iss


# ═══════ Пограничные случаи ═══════

class TestDocxEdgeCases:
    def test_empty_document(self, tmp_path):
        doc = Document()
        path = tmp_path / "empty.docx"
        doc.save(path)
        r = _check(path)
        assert r["file_type"] == "docx"

    def test_broken_file(self, tmp_path):
        path = tmp_path / "broken.docx"
        path.write_bytes(b"not a real docx")
        r = _check(path)
        assert "error" in r
        assert r["total_issues"] == 0

    def test_unicode_content(self, tmp_path):
        path = _make_doc(tmp_path, text="Текст с ёжиком 🦔 и формулой: α + β = γ")
        r = _check(path)
        assert r["file_type"] == "docx"

    def test_many_paragraphs(self, tmp_path):
        texts = [f"Абзац номер {i} содержит текст." for i in range(50)]
        path = _make_doc(tmp_path, paragraphs=texts)
        r = _check(path)
        # 50 основных + 2 библиография = 52
        assert r["paragraphs_checked"] == 52

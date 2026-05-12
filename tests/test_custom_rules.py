"""Тесты для пользовательских правил проверки .docx (v6 фича)."""
import io
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm, Pt

from app.checkers.docx_checker import (
    _get_rules,
    PRESETS,
    check_docx_document,
)
from app.checkers.docx_fixer import autofix_docx


# ─── Хелперы ───

def _make_doc_with_font(path: Path, font_name: str, font_size_pt: float) -> None:
    """Создаёт минимальный документ с заданным шрифтом и размером."""
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    p = doc.add_paragraph()
    run = p.add_run("Тестовый абзац с определённым шрифтом.")
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.25)
    doc.save(path)


# ─── Тесты _get_rules ───

def test_get_rules_default_returns_agu():
    """Без аргументов — дефолтные правила АГУ."""
    rules = _get_rules(None)
    assert rules["font_name"] == "Times New Roman"
    assert rules["font_size_pt"] == 14.0
    assert rules["line_spacing"] == 1.5
    assert rules["first_line_indent_cm"] == 1.25
    assert rules["margins_cm"]["left"] == 3.0


def test_get_rules_preset_agu():
    """Пресет 'agu' возвращает те же значения, что дефолт."""
    rules = _get_rules({"preset": "agu"})
    default = _get_rules(None)
    assert rules["font_name"] == default["font_name"]
    assert rules["font_size_pt"] == default["font_size_pt"]
    assert rules["margins_cm"] == default["margins_cm"]


def test_get_rules_custom_overrides():
    """Кастомные значения перезаписывают дефолтные."""
    rules = _get_rules({
        "font_name": "Arial",
        "font_size_pt": 12,
        "margins_cm": {"left": 2.5},
    })
    assert rules["font_name"] == "Arial"
    assert rules["font_size_pt"] == 12.0
    # Незатронутые поля остаются дефолтными
    assert rules["line_spacing"] == 1.5
    assert rules["margins_cm"]["right"] == 1.5  # не изменено
    assert rules["margins_cm"]["left"] == 2.5   # изменено


def test_get_rules_unknown_preset_falls_back():
    """Неизвестный пресет использует дефолтные значения."""
    rules = _get_rules({"preset": "nonexistent_university"})
    default = _get_rules(None)
    assert rules["font_name"] == default["font_name"]


def test_get_rules_string_to_float():
    """Числовые строки корректно конвертируются в float."""
    rules = _get_rules({"font_size_pt": "12", "line_spacing": "2.0"})
    assert rules["font_size_pt"] == 12.0
    assert rules["line_spacing"] == 2.0


# ─── Тесты PRESETS ───

def test_presets_contain_agu():
    """Пресет АГУ должен существовать."""
    assert "agu" in PRESETS
    assert PRESETS["agu"]["font_name"] == "Times New Roman"


def test_all_presets_have_required_keys():
    """Все пресеты должны содержать одинаковый набор ключей."""
    required = {"font_name", "font_size_pt", "line_spacing",
                "first_line_indent_cm", "margins_cm", "label"}
    for name, preset in PRESETS.items():
        assert required.issubset(preset.keys()), (
            f"Пресет '{name}' не содержит ключи: {required - set(preset.keys())}"
        )


# ─── Тесты проверки с кастомными правилами ───

def test_custom_rules_font_12_passes_with_12pt_doc(tmp_path):
    """Документ с 12pt шрифтом проходит проверку при правиле font_size_pt=12."""
    path = tmp_path / "doc12.docx"
    _make_doc_with_font(path, "Times New Roman", 12.0)

    # С дефолтными правилами (14pt) — ошибка
    report_default = check_docx_document(path, "doc12.docx", None)
    size_issues = [i for i in report_default["issues"] if i["code"] == "FONT_SIZE_MISMATCH"]
    assert len(size_issues) > 0, "Должна быть ошибка кегля при дефолтных правилах"

    # С кастомными правилами (12pt) — всё ОК
    report_custom = check_docx_document(path, "doc12.docx", {"font_size_pt": 12.0})
    size_issues = [i for i in report_custom["issues"] if i["code"] == "FONT_SIZE_MISMATCH"]
    assert len(size_issues) == 0, "Не должно быть ошибки кегля при правиле 12pt"


def test_custom_rules_arial_passes_with_arial_doc(tmp_path):
    """Документ с Arial проходит проверку при правиле font_name=Arial."""
    path = tmp_path / "arial.docx"
    _make_doc_with_font(path, "Arial", 14.0)

    # С дефолтными правилами — ошибка шрифта
    report_default = check_docx_document(path, "arial.docx", None)
    font_issues = [i for i in report_default["issues"] if i["code"] == "FONT_MISMATCH"]
    assert len(font_issues) > 0

    # С кастомными правилами (Arial) — ОК
    report_custom = check_docx_document(path, "arial.docx", {"font_name": "Arial"})
    font_issues = [i for i in report_custom["issues"] if i["code"] == "FONT_MISMATCH"]
    assert len(font_issues) == 0


# ─── Тесты автоисправления с кастомными правилами ───

def test_autofix_uses_custom_rules(tmp_path):
    """autofix с кастомными правилами (Arial, 12pt) применяет их."""
    path = tmp_path / "bad.docx"
    _make_doc_with_font(path, "Verdana", 10.0)

    custom = {"font_name": "Arial", "font_size_pt": 12}
    fixed_bytes = autofix_docx(path, custom)
    doc = Document(io.BytesIO(fixed_bytes))

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            for run in paragraph.runs:
                assert run.font.name == "Arial", f"Шрифт должен быть Arial, а не {run.font.name}"
                assert abs(run.font.size.pt - 12.0) < 0.5, f"Кегль должен быть 12, а не {run.font.size.pt}"


def test_autofix_custom_margins(tmp_path):
    """autofix с кастомными полями применяет их."""
    path = tmp_path / "bad.docx"
    _make_doc_with_font(path, "Times New Roman", 14.0)

    custom = {"margins_cm": {"left": 2.5, "right": 1.0}}
    fixed_bytes = autofix_docx(path, custom)
    doc = Document(io.BytesIO(fixed_bytes))
    sec = doc.sections[0]
    assert abs(sec.left_margin.cm - 2.5) < 0.1
    assert abs(sec.right_margin.cm - 1.0) < 0.1

"""Тесты для расширенных проверок: подписи и список литературы."""
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm, Pt

from app.checkers.docx_extras import (
    check_bibliography,
    check_figure_captions,
    check_table_captions,
)


def _make_doc_with_table(path: Path, caption: str | None = None) -> None:
    """Создаёт документ с таблицей (и опционально — подписью перед ней)."""
    doc = Document()
    if caption:
        doc.add_paragraph(caption)
    doc.add_table(rows=2, cols=2)
    doc.save(path)


def _make_doc_with_bibliography(path: Path, header: str, entries: list[str]) -> None:
    """Создаёт документ с разделом списка литературы."""
    doc = Document()
    doc.add_paragraph("Текст работы.")
    doc.add_paragraph(header)
    for entry in entries:
        doc.add_paragraph(entry)
    doc.save(path)


def test_table_no_caption(tmp_path):
    """Таблица без подписи → TABLE_NO_CAPTION."""
    path = tmp_path / "no_caption.docx"
    _make_doc_with_table(path, caption=None)
    doc = Document(path)
    issues = check_table_captions(doc)
    codes = {i["code"] for i in issues}
    assert "TABLE_NO_CAPTION" in codes


def test_table_correct_caption(tmp_path):
    """Правильная подпись «Таблица 1 — Название» не даёт замечаний."""
    path = tmp_path / "good_caption.docx"
    _make_doc_with_table(path, caption="Таблица 1 — Результаты эксперимента")
    doc = Document(path)
    issues = check_table_captions(doc)
    assert len(issues) == 0


def test_table_wrong_dash(tmp_path):
    """Дефис вместо длинного тире → TABLE_DASH_WRONG."""
    path = tmp_path / "bad_dash.docx"
    _make_doc_with_table(path, caption="Таблица 1 - Результаты эксперимента")
    doc = Document(path)
    issues = check_table_captions(doc)
    codes = {i["code"] for i in issues}
    assert "TABLE_DASH_WRONG" in codes


def test_bibliography_missing(tmp_path):
    """Документ без списка литературы → BIBLIOGRAPHY_MISSING."""
    path = tmp_path / "no_bib.docx"
    doc = Document()
    doc.add_paragraph("Просто текст без списка литературы.")
    doc.save(path)
    doc = Document(path)
    issues = check_bibliography(doc)
    codes = {i["code"] for i in issues}
    assert "BIBLIOGRAPHY_MISSING" in codes


def test_bibliography_correct(tmp_path):
    """Правильный список литературы не даёт замечаний."""
    path = tmp_path / "good_bib.docx"
    _make_doc_with_bibliography(
        path,
        header="Список литературы",
        entries=[
            "1. Иванов И. И. Основы программирования. — М.: Наука, 2020.",
            "2. Петров П. П. Базы данных. — СПб.: Питер, 2019.",
        ],
    )
    doc = Document(path)
    issues = check_bibliography(doc)
    assert len(issues) == 0


def test_bibliography_empty(tmp_path):
    """Заголовок есть, но записей нет → BIBLIOGRAPHY_EMPTY."""
    path = tmp_path / "empty_bib.docx"
    _make_doc_with_bibliography(
        path,
        header="Список использованных источников",
        entries=[],
    )
    doc = Document(path)
    issues = check_bibliography(doc)
    codes = {i["code"] for i in issues}
    assert "BIBLIOGRAPHY_EMPTY" in codes


def test_figure_caption_wrong_format(tmp_path):
    """«Рис. 1 - ...» вместо «Рисунок 1 — ...» → FIGURE_CAPTION_FORMAT."""
    path = tmp_path / "bad_fig.docx"
    doc = Document()
    doc.add_paragraph("Рис. 1 - Схема системы")
    doc.save(path)
    doc = Document(path)
    issues = check_figure_captions(doc)
    codes = {i["code"] for i in issues}
    assert "FIGURE_CAPTION_FORMAT" in codes

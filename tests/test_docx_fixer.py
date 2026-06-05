"""Тесты для модуля автоисправления .docx."""
import io
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm, Pt

from app.checkers.docx_fixer import autofix_docx
from app.checkers.docx_checker import check_docx_document


def _make_bad_doc(path: Path) -> None:
    """Создаёт документ с нарушениями."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    p = doc.add_paragraph()
    run = p.add_run("Текст с нарушениями оформления.")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.0
    doc.save(path)


def test_autofix_returns_bytes(tmp_path):
    """autofix_docx должен вернуть байты — валидный .docx."""
    path = tmp_path / "bad.docx"
    _make_bad_doc(path)
    result = autofix_docx(path)
    assert isinstance(result, bytes)
    assert len(result) > 100  # не пустой
    # Можно открыть обратно как документ
    doc = Document(io.BytesIO(result))
    assert len(doc.paragraphs) > 0


def test_autofix_fixes_margins(tmp_path):
    """После autofix поля должны быть 3/1.5/2/2."""
    path = tmp_path / "bad.docx"
    _make_bad_doc(path)
    result = autofix_docx(path)
    doc = Document(io.BytesIO(result))
    section = doc.sections[0]
    assert abs(section.left_margin.cm - 3.0) < 0.1
    assert abs(section.right_margin.cm - 1.5) < 0.1
    assert abs(section.top_margin.cm - 2.0) < 0.1
    assert abs(section.bottom_margin.cm - 2.0) < 0.1


def test_autofix_fixes_font(tmp_path):
    """После autofix шрифт должен быть Times New Roman 14."""
    path = tmp_path / "bad.docx"
    _make_bad_doc(path)
    result = autofix_docx(path)
    doc = Document(io.BytesIO(result))
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            for run in paragraph.runs:
                assert run.font.name == "Times New Roman"
                assert abs(run.font.size.pt - 14.0) < 0.5


def test_autofix_reduces_issues(tmp_path):
    """После autofix количество замечаний должно уменьшиться."""
    path = tmp_path / "bad.docx"
    _make_bad_doc(path)

    # Считаем замечания до исправления
    report_before = check_docx_document(path, "bad.docx")
    issues_before = report_before["total_issues"]
    assert issues_before > 0  # файл действительно плохой

    # Исправляем и проверяем снова
    fixed_bytes = autofix_docx(path)
    fixed_path = tmp_path / "fixed.docx"
    fixed_path.write_bytes(fixed_bytes)
    report_after = check_docx_document(fixed_path, "fixed.docx")
    issues_after = report_after["total_issues"]

    assert issues_after < issues_before, (
        f"До: {issues_before} замечаний, после: {issues_after}. "
        f"Autofix должен был уменьшить число ошибок."
    )

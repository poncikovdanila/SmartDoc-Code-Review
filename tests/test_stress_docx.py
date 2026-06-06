"""Стресс-тесты: документы с хаотичным форматированием.

Проверяют, что чекер корректно обнаруживает проблемы в документах,
максимально приближенных к реальным студенческим работам:
разные шрифты, цвета, размеры, отступы, интервалы в одном документе.
"""
import io
import re
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.checkers.docx_checker import check_docx_document
from app.checkers.docx_fixer import autofix_docx


# ═══════ Хелперы ═══════

def _save(doc, tmp_path, name="test.docx"):
    path = tmp_path / name
    doc.save(path)
    return path


def _check(path, rules=None):
    return check_docx_document(path, path.name, rules)


def _codes(report):
    return {i["code"] for i in report["issues"]}


def _count(report, code):
    return sum(1 for i in report["issues"] if i["code"] == code)


def _make_base_doc():
    """Базовый документ с правильными полями."""
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Cm(3); sec.right_margin = Cm(1.5)
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
    return doc


def _add_paragraph(doc, text, font="Times New Roman", size=14,
                   spacing=1.5, indent=1.25, color=None,
                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
                   italic=False, underline=False):
    """Добавляет абзац с полным контролем форматирования."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.first_line_indent = Cm(indent)
    p.alignment = alignment
    return p


# ═══════ 1. Разноцветный текст ═══════

class TestColoredText:
    """Документ, где каждый абзац разного цвета."""

    def test_all_colors_detected(self, tmp_path):
        doc = _make_base_doc()
        colors = [
            (255, 0, 0),    # красный
            (0, 0, 255),    # синий
            (0, 128, 0),    # зелёный
            (128, 0, 128),  # фиолетовый
            (255, 165, 0),  # оранжевый
        ]
        for i, color in enumerate(colors):
            _add_paragraph(doc, f"Абзац {i+1} с цветным текстом.", color=color)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "COLORED_TEXT" in _codes(r)

    def test_mixed_colors_in_one_paragraph(self, tmp_path):
        doc = _make_base_doc()
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Один run чёрный, другой красный, третий синий
        r1 = p.add_run("Начало предложения ")
        r1.font.name = "Times New Roman"; r1.font.size = Pt(14)
        r1.font.color.rgb = RGBColor(0, 0, 0)
        r2 = p.add_run("красное слово ")
        r2.font.name = "Times New Roman"; r2.font.size = Pt(14)
        r2.font.color.rgb = RGBColor(255, 0, 0)
        r3 = p.add_run("и синее слово.")
        r3.font.name = "Times New Roman"; r3.font.size = Pt(14)
        r3.font.color.rgb = RGBColor(0, 0, 255)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "COLORED_TEXT" in _codes(r)

    def test_all_black_no_color_issue(self, tmp_path):
        doc = _make_base_doc()
        for i in range(5):
            _add_paragraph(doc, f"Абзац {i+1} полностью чёрный.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "COLORED_TEXT" not in _codes(r)


# ═══════ 2. Разные шрифты в одном документе ═══════

class TestMixedFonts:
    """Документ с разными шрифтами в разных абзацах."""

    def test_multiple_wrong_fonts(self, tmp_path):
        doc = _make_base_doc()
        fonts = ["Arial", "Courier New", "Verdana", "Calibri", "Comic Sans MS"]
        for font in fonts:
            _add_paragraph(doc, f"Текст шрифтом {font}.", font=font)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "FONT_MISMATCH") >= 5

    def test_mixed_fonts_in_one_paragraph(self, tmp_path):
        doc = _make_base_doc()
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for font in ["Times New Roman", "Arial", "Courier New"]:
            run = p.add_run(f"Текст {font}. ")
            run.font.name = font
            run.font.size = Pt(14)
        path = _save(doc, tmp_path)
        r = _check(path)
        # Первый run — TNR, значит FONT_MISMATCH не сработает (берёт первый run)
        # Но это нормальное поведение — чекер берёт первый run

    def test_correct_font_no_issues(self, tmp_path):
        doc = _make_base_doc()
        for _ in range(5):
            _add_paragraph(doc, "Текст шрифтом Times New Roman.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "FONT_MISMATCH") == 0


# ═══════ 3. Разные размеры шрифта ═══════

class TestMixedFontSizes:
    """Документ с разными кеглями."""

    def test_various_sizes(self, tmp_path):
        doc = _make_base_doc()
        for size in [8, 10, 11, 12, 16, 18, 20, 24]:
            _add_paragraph(doc, f"Текст размером {size} пт.", size=size)
        path = _save(doc, tmp_path)
        r = _check(path)
        # Все кроме 14 — ошибка (допуск ±0.5)
        assert _count(r, "FONT_SIZE_MISMATCH") == 8

    def test_almost_correct_size(self, tmp_path):
        doc = _make_base_doc()
        _add_paragraph(doc, "Текст размером 14.0 пт.", size=14.0)
        _add_paragraph(doc, "Текст размером 13.8 пт.", size=13.8)
        _add_paragraph(doc, "Текст размером 14.2 пт.", size=14.2)
        path = _save(doc, tmp_path)
        r = _check(path)
        # 13.8 и 14.2 в пределах допуска ±0.5
        assert _count(r, "FONT_SIZE_MISMATCH") == 0


# ═══════ 4. Разные отступы ═══════

class TestMixedIndents:
    """Документ с разными отступами первой строки."""

    def test_various_indents(self, tmp_path):
        doc = _make_base_doc()
        for indent in [0, 0.5, 1.0, 1.25, 2.0, 2.5, 3.0]:
            _add_paragraph(doc, f"Отступ {indent} см.", indent=indent)
        path = _save(doc, tmp_path)
        r = _check(path)
        # Только 1.25 — правильный (допуск ±0.15)
        correct_count = sum(1 for ind in [0, 0.5, 1.0, 1.25, 2.0, 2.5, 3.0]
                           if abs(ind - 1.25) <= 0.15)
        wrong_count = 7 - correct_count
        assert _count(r, "INDENT_MISMATCH") >= wrong_count

    def test_no_indent(self, tmp_path):
        doc = _make_base_doc()
        for _ in range(5):
            _add_paragraph(doc, "Текст без отступа.", indent=0)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "INDENT_MISMATCH") == 5


# ═══════ 5. Разные межстрочные интервалы ═══════

class TestMixedSpacing:
    """Документ с разными интервалами."""

    def test_various_spacings(self, tmp_path):
        doc = _make_base_doc()
        for sp in [1.0, 1.15, 1.5, 2.0, 2.5, 3.0]:
            _add_paragraph(doc, f"Интервал {sp}.", spacing=sp)
        path = _save(doc, tmp_path)
        r = _check(path)
        # Только 1.5 правильный
        assert _count(r, "LINE_SPACING_MISMATCH") >= 5

    def test_correct_spacing(self, tmp_path):
        doc = _make_base_doc()
        for _ in range(5):
            _add_paragraph(doc, "Правильный интервал.", spacing=1.5)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "LINE_SPACING_MISMATCH") == 0


# ═══════ 6. Разные поля страницы ═══════

class TestMixedMargins:
    """Документ с неправильными полями."""

    def test_all_margins_wrong(self, tmp_path):
        doc = Document()
        sec = doc.sections[0]
        sec.left_margin = Cm(1.0)
        sec.right_margin = Cm(3.0)
        sec.top_margin = Cm(1.0)
        sec.bottom_margin = Cm(1.0)
        _add_paragraph(doc, "Текст с неправильными полями.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "MARGIN_MISMATCH") == 4

    def test_two_sections_different_margins(self, tmp_path):
        doc = Document()
        # Секция 1 — правильные поля
        sec1 = doc.sections[0]
        sec1.left_margin = Cm(3); sec1.right_margin = Cm(1.5)
        sec1.top_margin = Cm(2); sec1.bottom_margin = Cm(2)
        _add_paragraph(doc, "Секция 1 — правильные поля.")
        # Секция 2 — неправильные
        doc.add_section()
        sec2 = doc.sections[1]
        sec2.left_margin = Cm(2); sec2.right_margin = Cm(2)
        sec2.top_margin = Cm(3); sec2.bottom_margin = Cm(3)
        _add_paragraph(doc, "Секция 2 — неправильные поля.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "MARGIN_MISMATCH") >= 4


# ═══════ 7. Выравнивание текста ═══════

class TestMixedAlignment:
    """Документ с разным выравниванием."""

    def test_left_aligned(self, tmp_path):
        doc = _make_base_doc()
        for _ in range(3):
            _add_paragraph(doc, "Текст, выровненный по левому краю, что неправильно для основного текста.",
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "ALIGN_NOT_JUSTIFY" in _codes(r)

    def test_center_aligned(self, tmp_path):
        doc = _make_base_doc()
        _add_paragraph(doc, "Текст по центру — тоже неправильно для основного текста, если он длинный.",
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "ALIGN_NOT_JUSTIFY" in _codes(r)

    def test_right_aligned(self, tmp_path):
        doc = _make_base_doc()
        _add_paragraph(doc, "Текст по правому краю — тоже неправильно для основного текста работы.",
                      alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "ALIGN_NOT_JUSTIFY" in _codes(r)

    def test_justified_no_issue(self, tmp_path):
        doc = _make_base_doc()
        for _ in range(3):
            _add_paragraph(doc, "Текст по ширине — правильное выравнивание для основного текста.",
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "ALIGN_NOT_JUSTIFY" not in _codes(r)


# ═══════ 8. Заголовки ═══════

class TestHeadings:
    """Документ с проблемными заголовками."""

    def test_heading_with_dot(self, tmp_path):
        doc = _make_base_doc()
        h = doc.add_heading("Введение.", level=1)
        for run in h.runs:
            run.bold = True
        _add_paragraph(doc, "Текст после заголовка.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "HEADING_ENDS_WITH_DOT" in _codes(r)

    def test_heading_not_bold(self, tmp_path):
        doc = _make_base_doc()
        h = doc.add_heading("Введение", level=1)
        for run in h.runs:
            run.bold = False
        _add_paragraph(doc, "Текст после заголовка.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "HEADING_NOT_BOLD" in _codes(r)

    def test_correct_heading(self, tmp_path):
        doc = _make_base_doc()
        h = doc.add_heading("Введение", level=1)
        for run in h.runs:
            run.bold = True
        _add_paragraph(doc, "Текст после заголовка.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "HEADING_ENDS_WITH_DOT" not in _codes(r)


# ═══════ 9. Пробелы ═══════

class TestSpaces:
    """Документ с проблемами пробелов."""

    def test_double_spaces(self, tmp_path):
        doc = _make_base_doc()
        _add_paragraph(doc, "Текст  с  двойными  пробелами  везде.")
        _add_paragraph(doc, "Ещё   тройные   пробелы   тут.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "DOUBLE_SPACES" in _codes(r)

    def test_space_before_comma(self, tmp_path):
        doc = _make_base_doc()
        _add_paragraph(doc, "Текст , с пробелом перед запятой .")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "SPACE_BEFORE_PUNCT" in _codes(r)

    def test_clean_text_no_space_issues(self, tmp_path):
        doc = _make_base_doc()
        _add_paragraph(doc, "Правильный текст, без лишних пробелов.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "DOUBLE_SPACES" not in _codes(r)
        assert "SPACE_BEFORE_PUNCT" not in _codes(r)


# ═══════ 10. Пустые строки ═══════

class TestBlankLines:
    """Документ с лишними пустыми строками."""

    def test_many_blank_lines(self, tmp_path):
        doc = _make_base_doc()
        _add_paragraph(doc, "Абзац перед пустыми строками.")
        for _ in range(5):
            doc.add_paragraph()
        _add_paragraph(doc, "Абзац после пустых строк.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "EXTRA_BLANK_LINES" in _codes(r)

    def test_one_blank_line_ok(self, tmp_path):
        doc = _make_base_doc()
        _add_paragraph(doc, "Первый абзац.")
        doc.add_paragraph()
        _add_paragraph(doc, "Второй абзац.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "EXTRA_BLANK_LINES" not in _codes(r)


# ═══════ 11. Гиперссылки ═══════

class TestHyperlinks:
    """Документ с гиперссылками."""

    def test_hyperlink_detected(self, tmp_path):
        doc = _make_base_doc()
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(1.25)
        # Создаём гиперссылку через XML
        r = p.add_run("Обычный текст. ")
        r.font.name = "Times New Roman"; r.font.size = Pt(14)
        # Добавляем гиперссылку
        hyperlink = p._element.makeelement(qn("w:hyperlink"), {
            qn("r:id"): "rId1",
        })
        new_run = p._element.makeelement(qn("w:r"), {})
        rpr = new_run.makeelement(qn("w:rPr"), {})
        new_run.append(rpr)
        t = new_run.makeelement(qn("w:t"), {})
        t.text = "ссылка"
        new_run.append(t)
        hyperlink.append(new_run)
        p._element.append(hyperlink)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "HYPERLINKS_FOUND" in _codes(r)


# ═══════ 12. Таблицы ═══════

class TestTables:
    """Документ с таблицами разного форматирования."""

    def test_wrong_font_in_table(self, tmp_path):
        doc = _make_base_doc()
        table = doc.add_table(rows=2, cols=2)
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    run = p.add_run("Текст Arial")
                    run.font.name = "Arial"
                    run.font.size = Pt(12)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "TABLE_FONT_MISMATCH" in _codes(r)

    def test_correct_font_in_table(self, tmp_path):
        doc = _make_base_doc()
        table = doc.add_table(rows=2, cols=2)
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    run = p.add_run("Текст TNR")
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(14)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert "TABLE_FONT_MISMATCH" not in _codes(r)


# ═══════ 13. Полный хаос ═══════

class TestFullChaos:
    """Документ, где ВСЁ неправильно — максимальный стресс-тест."""

    def test_everything_wrong(self, tmp_path):
        doc = Document()
        sec = doc.sections[0]
        sec.left_margin = Cm(1); sec.right_margin = Cm(1)
        sec.top_margin = Cm(1); sec.bottom_margin = Cm(1)

        # Заголовок с точкой, не жирный
        h = doc.add_heading("Глава 1.", level=1)
        for run in h.runs:
            run.bold = False

        # Абзацы с разными ошибками
        _add_paragraph(doc, "Текст Arial красный размер 20.", font="Arial", size=20,
                      color=(255,0,0), spacing=2.0, indent=0,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT)
        _add_paragraph(doc, "Текст  с  двойными  пробелами , запятая.", font="Courier New",
                      size=10, color=(0,0,255), spacing=1.0, indent=3.0,
                      alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        _add_paragraph(doc, "Verdana зелёный подчёркнутый.", font="Verdana", size=16,
                      color=(0,128,0), spacing=3.0, indent=0.5,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, underline=True)

        # Пустые строки
        for _ in range(4):
            doc.add_paragraph()

        _add_paragraph(doc, "Ещё абзац после пустых строк.", font="Calibri", size=8,
                      color=(128,0,128), spacing=1.15, indent=2.5)

        # Таблица с Arial
        table = doc.add_table(rows=2, cols=2)
        for row in table.rows:
            for cell in row.cells:
                run = cell.paragraphs[0].add_run("Arial в таблице")
                run.font.name = "Arial"

        path = _save(doc, tmp_path)
        r = _check(path)

        expected_codes = {
            "MARGIN_MISMATCH", "FONT_MISMATCH", "FONT_SIZE_MISMATCH",
            "LINE_SPACING_MISMATCH", "INDENT_MISMATCH", "ALIGN_NOT_JUSTIFY",
            "COLORED_TEXT", "DOUBLE_SPACES", "SPACE_BEFORE_PUNCT",
            "EXTRA_BLANK_LINES", "HEADING_ENDS_WITH_DOT", "HEADING_NOT_BOLD",
            "TABLE_FONT_MISMATCH",
        }
        found = _codes(r)
        for code in expected_codes:
            assert code in found, f"Ожидался код {code}, но не найден. Найдены: {found}"

    def test_chaos_autofix_reduces_issues(self, tmp_path):
        """После autofix количество ошибок должно уменьшиться."""
        doc = Document()
        sec = doc.sections[0]
        sec.left_margin = Cm(1); sec.right_margin = Cm(1)
        sec.top_margin = Cm(1); sec.bottom_margin = Cm(1)
        _add_paragraph(doc, "Arial  красный  двойные  пробелы .",
                      font="Arial", size=20, color=(255,0,0),
                      spacing=2.0, indent=0, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        _add_paragraph(doc, "Courier синий.", font="Courier New", size=10,
                      color=(0,0,255), spacing=1.0, indent=3.0)
        path = _save(doc, tmp_path)

        before = _check(path)
        fixed_bytes = autofix_docx(path)
        fixed_path = tmp_path / "fixed.docx"
        fixed_path.write_bytes(fixed_bytes)
        after = _check(fixed_path)

        assert after["total_issues"] < before["total_issues"], \
            f"Autofix не уменьшил ошибки: {before['total_issues']} → {after['total_issues']}"

    def test_chaos_verdict_improves_after_fix(self, tmp_path):
        """Verdict должен улучшиться после autofix."""
        doc = Document()
        sec = doc.sections[0]
        sec.left_margin = Cm(1); sec.right_margin = Cm(3)
        sec.top_margin = Cm(1); sec.bottom_margin = Cm(3)
        for i in range(10):
            _add_paragraph(doc, f"Абзац {i} с ошибками форматирования.",
                          font="Verdana", size=18, spacing=2.5, indent=0,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
        path = _save(doc, tmp_path)

        before = _check(path)
        fixed_bytes = autofix_docx(path)
        fixed_path = tmp_path / "fixed.docx"
        fixed_path.write_bytes(fixed_bytes)
        after = _check(fixed_path)

        assert after["total_issues"] < before["total_issues"], \
            f"Autofix не уменьшил ошибки: {before['total_issues']} → {after['total_issues']}"


# ═══════ 14. Автоисправление конкретных проблем ═══════

class TestAutofixSpecific:
    """Проверяет, что autofix исправляет конкретные виды ошибок."""

    def test_fixes_colored_text(self, tmp_path):
        doc = _make_base_doc()
        _add_paragraph(doc, "Красный текст.", color=(255, 0, 0))
        _add_paragraph(doc, "Синий текст.", color=(0, 0, 255))
        path = _save(doc, tmp_path)
        fixed = autofix_docx(path)
        fdoc = Document(io.BytesIO(fixed))
        for p in fdoc.paragraphs:
            for run in p.runs:
                if run.text.strip() and run.font.color.rgb:
                    assert run.font.color.rgb == RGBColor(0, 0, 0), \
                        f"Цвет не исправлен: {run.font.color.rgb}"

    def test_fixes_double_spaces(self, tmp_path):
        doc = _make_base_doc()
        _add_paragraph(doc, "Текст  с  двойными  пробелами.")
        path = _save(doc, tmp_path)
        fixed = autofix_docx(path)
        fdoc = Document(io.BytesIO(fixed))
        for p in fdoc.paragraphs:
            assert "  " not in p.text, f"Двойные пробелы не удалены: '{p.text}'"

    def test_fixes_space_before_punct(self, tmp_path):
        doc = _make_base_doc()
        _add_paragraph(doc, "Текст , с пробелом . перед знаками !")
        path = _save(doc, tmp_path)
        fixed = autofix_docx(path)
        fdoc = Document(io.BytesIO(fixed))
        for p in fdoc.paragraphs:
            text = p.text
            if text.strip():
                assert not re.search(r" +[,.]", text), \
                    f"Пробелы перед знаками не удалены: '{text}'"

    def test_fixes_alignment(self, tmp_path):
        doc = _make_base_doc()
        _add_paragraph(doc, "Текст по левому краю.", alignment=WD_ALIGN_PARAGRAPH.LEFT)
        path = _save(doc, tmp_path)
        fixed = autofix_docx(path)
        fdoc = Document(io.BytesIO(fixed))
        for p in fdoc.paragraphs:
            if p.text.strip() and len(p.text.strip()) > 10:
                assert p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY, \
                    f"Выравнивание не исправлено: {p.alignment}"

    def test_fixes_heading_dot(self, tmp_path):
        doc = _make_base_doc()
        h = doc.add_heading("Введение.", level=1)
        for run in h.runs:
            run.bold = True
        path = _save(doc, tmp_path)
        fixed = autofix_docx(path)
        fdoc = Document(io.BytesIO(fixed))
        for p in fdoc.paragraphs:
            style = (p.style.name or "").lower()
            if "heading" in style and p.text.strip():
                assert not p.text.strip().endswith("."), \
                    f"Точка в заголовке не удалена: '{p.text}'"

    def test_fixes_heading_bold(self, tmp_path):
        doc = _make_base_doc()
        h = doc.add_heading("Введение", level=1)
        for run in h.runs:
            run.bold = False
        path = _save(doc, tmp_path)
        fixed = autofix_docx(path)
        fdoc = Document(io.BytesIO(fixed))
        for p in fdoc.paragraphs:
            style = (p.style.name or "").lower()
            if "heading" in style and p.text.strip():
                for run in p.runs:
                    if run.text.strip():
                        assert run.bold, f"Заголовок не сделан жирным"

    def test_fixes_wrong_font(self, tmp_path):
        doc = _make_base_doc()
        _add_paragraph(doc, "Arial текст.", font="Arial")
        _add_paragraph(doc, "Verdana текст.", font="Verdana")
        path = _save(doc, tmp_path)
        fixed = autofix_docx(path)
        fdoc = Document(io.BytesIO(fixed))
        for p in fdoc.paragraphs:
            for run in p.runs:
                if run.text.strip():
                    assert run.font.name == "Times New Roman", \
                        f"Шрифт не исправлен: {run.font.name}"

    def test_fixes_wrong_size(self, tmp_path):
        doc = _make_base_doc()
        _add_paragraph(doc, "Текст 20 пт.", size=20)
        _add_paragraph(doc, "Текст 10 пт.", size=10)
        path = _save(doc, tmp_path)
        fixed = autofix_docx(path)
        fdoc = Document(io.BytesIO(fixed))
        for p in fdoc.paragraphs:
            for run in p.runs:
                if run.text.strip() and run.font.size:
                    assert abs(run.font.size.pt - 14) < 1, \
                        f"Размер не исправлен: {run.font.size.pt}"


# ═══════ 15. Проверка с кастомными правилами ═══════

class TestChaosWithCustomRules:
    """Стресс-тесты с кастомными правилами."""

    def test_arial_12pt_accepted_with_custom(self, tmp_path):
        doc = _make_base_doc()
        for _ in range(5):
            _add_paragraph(doc, "Текст Arial 12pt.", font="Arial", size=12)
        path = _save(doc, tmp_path)

        # Default rules — ошибки
        r_default = _check(path)
        assert _count(r_default, "FONT_MISMATCH") > 0
        assert _count(r_default, "FONT_SIZE_MISMATCH") > 0

        # Custom rules — без ошибок шрифта
        r_custom = _check(path, {"font_name": "Arial", "font_size_pt": 12})
        assert _count(r_custom, "FONT_MISMATCH") == 0
        assert _count(r_custom, "FONT_SIZE_MISMATCH") == 0

    def test_disabled_checks(self, tmp_path):
        doc = _make_base_doc()
        h = doc.add_heading("Заголовок.", level=1)
        for run in h.runs:
            run.bold = False
        _add_paragraph(doc, "Текст  с  пробелами .", color=(255, 0, 0))
        path = _save(doc, tmp_path)

        # All checks on
        r_all = _check(path)
        all_count = r_all["total_issues"]

        # Disable some checks
        r_less = _check(path, {"checks": {
            "headings": False, "spaces": False, "textColor": False
        }})
        assert r_less["total_issues"] < all_count

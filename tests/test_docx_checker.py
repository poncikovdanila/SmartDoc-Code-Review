"""Тесты для модуля нормоконтроля .docx."""
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm, Pt

from app.checkers.docx_checker import check_docx_document


def _make_compliant_doc(path: Path) -> None:
    """Создаёт документ, полностью соответствующий ГОСТ/АГУ."""
    doc = Document()

    # Поля страницы
    section = doc.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Стиль Normal
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    p = doc.add_paragraph("Это правильно оформленный абзац основного текста.")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.25)

    # Список литературы — нужен для расширенных проверок
    bib_header = doc.add_paragraph("Список литературы")
    bib_header.paragraph_format.line_spacing = 1.5
    bib_header.paragraph_format.first_line_indent = Cm(1.25)
    bib = doc.add_paragraph("1. Иванов И. И. Основы информатики. — М.: Наука, 2020.")
    bib.paragraph_format.line_spacing = 1.5
    bib.paragraph_format.first_line_indent = Cm(1.25)

    doc.save(path)


def _make_non_compliant_doc(path: Path) -> None:
    """Создаёт документ с грубыми нарушениями всех проверяемых параметров."""
    doc = Document()

    # Неправильные поля
    section = doc.sections[0]
    section.left_margin = Cm(2.0)  # должно быть 3
    section.right_margin = Cm(2.0)  # должно быть 1.5
    section.top_margin = Cm(2.5)  # должно быть 2
    section.bottom_margin = Cm(2.5)  # должно быть 2

    # Неправильный шрифт и размер на уровне run
    p = doc.add_paragraph()
    run = p.add_run("Текст с нарушениями оформления.")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.0  # должно быть 1.5
    p.paragraph_format.first_line_indent = Cm(0.5)  # должно быть 1.25

    doc.save(path)


def test_compliant_doc_has_no_issues(tmp_path):
    """Документ, соответствующий ГОСТ, не должен порождать замечаний."""
    path = tmp_path / "good.docx"
    _make_compliant_doc(path)

    report = check_docx_document(path, "good.docx")

    assert report["file_type"] == "docx"
    assert report["total_issues"] == 0, (
        f"Найдены неожиданные замечания: {report['issues']}"
    )


def test_non_compliant_doc_finds_all_violations(tmp_path):
    """Документ с нарушениями должен породить замечания всех видов."""
    path = tmp_path / "bad.docx"
    _make_non_compliant_doc(path)

    report = check_docx_document(path, "bad.docx")

    assert report["total_issues"] > 0
    codes = {issue["code"] for issue in report["issues"]}

    # Поля
    assert "MARGIN_MISMATCH" in codes
    # Шрифт
    assert "FONT_MISMATCH" in codes
    # Размер
    assert "FONT_SIZE_MISMATCH" in codes
    # Интервал
    assert "LINE_SPACING_MISMATCH" in codes
    # Отступ
    assert "INDENT_MISMATCH" in codes


def test_margin_mismatch_includes_expected_and_actual(tmp_path):
    """Замечание о полях должно содержать ожидаемое и фактическое значения."""
    path = tmp_path / "margins.docx"
    _make_non_compliant_doc(path)

    report = check_docx_document(path, "margins.docx")
    margin_issues = [i for i in report["issues"] if i["code"] == "MARGIN_MISMATCH"]
    assert len(margin_issues) > 0

    for issue in margin_issues:
        assert "expected" in issue
        assert "actual" in issue
        assert "см" in issue["expected"]


def test_paragraphs_checked_count(tmp_path):
    """Поле paragraphs_checked должно отражать количество проверенных абзацев."""
    path = tmp_path / "multi.docx"
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(14)

    for i in range(3):
        p = doc.add_paragraph(f"Содержательный абзац номер {i + 1}.")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(1.25)
    # Пустой абзац — не должен учитываться
    doc.add_paragraph("")

    doc.save(path)
    report = check_docx_document(path, "multi.docx")
    assert report["paragraphs_checked"] == 3


def test_invalid_file_returns_error(tmp_path):
    """Битый .docx должен возвращать поле error без падения."""
    path = tmp_path / "broken.docx"
    path.write_bytes(b"this is not a real docx")

    report = check_docx_document(path, "broken.docx")
    assert "error" in report
    assert report["total_issues"] == 0

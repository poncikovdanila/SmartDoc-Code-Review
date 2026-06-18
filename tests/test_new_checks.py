"""Тесты новых проверок (v8): интервалы, нумерация, иерархия,
ссылки, приложения, библиография + edge cases для кода.
"""
import io
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.checkers.docx_checker import check_docx_document
from app.checkers.code_checker import check_python_code


# ═══════ Хелперы ═══════

def _save(doc, tmp_path, name="test.docx"):
    path = tmp_path / name
    doc.save(path)
    return path


def _check(path, rules=None):
    return check_docx_document(path, path.name, rules)


def _codes(report):
    return {i["code"] for i in report["issues"]}


def _count(report, code):
    return sum(1 for i in report["issues"] if i["code"] == code)


def _make_base():
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Cm(3.5); sec.right_margin = Cm(1.0)
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    return doc


def _add_para(doc, text, font="Times New Roman", size=14,
              spacing=1.5, indent=1.25, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.first_line_indent = Cm(indent)
    p.alignment = alignment
    return p


def _add_heading(doc, text, level=1, bold=True):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.bold = bold
    return h


def _write_py(tmp_path, name, content):
    path = tmp_path / name
    path.write_text(content, encoding="utf-8")
    return path


# ═══════════════════════════════════════════
# 13. Интервалы до/после абзаца
# ═══════════════════════════════════════════

class TestParagraphSpacing:

    def test_no_spacing_is_clean(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст без интервалов до и после абзаца, всё корректно.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "PARA_SPACING_BEFORE") == 0
        assert _count(r, "PARA_SPACING_AFTER") == 0

    def test_large_spacing_before_detected(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        p = _add_para(doc, "Абзац с большим интервалом перед ним, это ошибка форматирования.")
        p.paragraph_format.space_before = Pt(12)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "PARA_SPACING_BEFORE") >= 1

    def test_large_spacing_after_detected(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        p = _add_para(doc, "Абзац с большим интервалом после него, это ошибка форматирования.")
        p.paragraph_format.space_after = Pt(12)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "PARA_SPACING_AFTER") >= 1

    def test_small_spacing_ok(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        p = _add_para(doc, "Абзац с минимальным интервалом, это допустимо вполне.")
        p.paragraph_format.space_before = Pt(2)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "PARA_SPACING_BEFORE") == 0

    def test_disabled_via_rules(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        p = _add_para(doc, "Абзац с интервалом, но проверка отключена через правила.")
        p.paragraph_format.space_before = Pt(24)
        path = _save(doc, tmp_path)
        r = _check(path, {"checks": {"paraSpacing": False}})
        assert _count(r, "PARA_SPACING_BEFORE") == 0


# ═══════════════════════════════════════════
# 14. Нумерация разделов
# ═══════════════════════════════════════════

class TestHeadingNumbering:

    def test_sequential_ok(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "1 Введение")
        _add_para(doc, "Текст первого раздела документа.")
        _add_heading(doc, "2 Основная часть")
        _add_para(doc, "Текст второго раздела.")
        _add_heading(doc, "3 Заключение")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "HEADING_NUM_GAP") == 0

    def test_gap_detected(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "1 Введение")
        _add_para(doc, "Текст.")
        _add_heading(doc, "3 Заключение")  # пропуск 2
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "HEADING_NUM_GAP") >= 1

    def test_disabled_via_rules(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "1 Введение")
        _add_heading(doc, "5 Заключение")
        path = _save(doc, tmp_path)
        r = _check(path, {"checks": {"headingNumbers": False}})
        assert _count(r, "HEADING_NUM_GAP") == 0


# ═══════════════════════════════════════════
# 15. Иерархия заголовков
# ═══════════════════════════════════════════

class TestHeadingHierarchy:

    def test_correct_hierarchy(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "Раздел 1", level=1)
        _add_heading(doc, "Подраздел 1.1", level=2)
        _add_para(doc, "Текст.")
        _add_heading(doc, "Раздел 2", level=1)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "HEADING_LEVEL_SKIP") == 0

    def test_level_skip_detected(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "Раздел 1", level=1)
        _add_heading(doc, "Пункт 1.1.1", level=3)  # перескок 1→3
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "HEADING_LEVEL_SKIP") >= 1

    def test_same_level_repeated_ok(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "Раздел 1", level=1)
        _add_heading(doc, "Раздел 2", level=1)
        _add_heading(doc, "Раздел 3", level=1)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "HEADING_LEVEL_SKIP") == 0


# ═══════════════════════════════════════════
# 16. Ссылки на рисунки и таблицы
# ═══════════════════════════════════════════

class TestCrossReferences:

    def test_figure_referenced_ok(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Как видно на рисунке 1, результат положительный.")
        p = doc.add_paragraph()
        run = p.add_run("Рисунок 1 — Результат эксперимента")
        run.font.size = Pt(14)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "FIGURE_NO_REFERENCE") == 0

    def test_figure_not_referenced(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст без упоминания рисунков.")
        p = doc.add_paragraph()
        run = p.add_run("Рисунок 1 — Схема")
        run.font.size = Pt(14)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "FIGURE_NO_REFERENCE") >= 1

    def test_table_referenced_ok(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "В таблице 1.1 представлены данные исследования.")
        p = doc.add_paragraph()
        run = p.add_run("Таблица 1.1 — Данные")
        run.font.size = Pt(14)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "TABLE_NO_REFERENCE") == 0

    def test_table_not_referenced(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст без упоминания таблиц вообще.")
        p = doc.add_paragraph()
        run = p.add_run("Таблица 2 — Результаты")
        run.font.size = Pt(14)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "TABLE_NO_REFERENCE") >= 1

    def test_figure_reference_in_parentheses(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Результат показан на графике (рис. 2.1).")
        p = doc.add_paragraph()
        run = p.add_run("Рисунок 2.1 — График")
        run.font.size = Pt(14)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "FIGURE_NO_REFERENCE") == 0

    def test_disabled_via_rules(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        p = doc.add_paragraph()
        run = p.add_run("Рисунок 1 — Без ссылки")
        run.font.size = Pt(14)
        path = _save(doc, tmp_path)
        r = _check(path, {"checks": {"crossReferences": False}})
        assert _count(r, "FIGURE_NO_REFERENCE") == 0


# ═══════════════════════════════════════════
# 17. Формат приложений
# ═══════════════════════════════════════════

class TestAppendixFormat:

    def test_correct_order(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ПРИЛОЖЕНИЕ А")
        _add_para(doc, "Содержание приложения А.")
        _add_heading(doc, "ПРИЛОЖЕНИЕ Б")
        _add_para(doc, "Содержание приложения Б.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "APPENDIX_ORDER") == 0

    def test_wrong_order(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ПРИЛОЖЕНИЕ Б")  # должно быть А
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "APPENDIX_ORDER") >= 1

    def test_single_appendix_ok(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ПРИЛОЖЕНИЕ А")
        _add_para(doc, "Единственное приложение.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "APPENDIX_ORDER") == 0

    def test_no_appendix_ok(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ЗАКЛЮЧЕНИЕ")
        _add_para(doc, "Текст заключения.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "APPENDIX_ORDER") == 0
        assert _count(r, "APPENDIX_FORMAT") == 0


# ═══════════════════════════════════════════
# 18. Формат библиографии
# ═══════════════════════════════════════════

class TestBibliographyFormat:

    def test_entries_with_year_ok(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
        _add_para(doc, "1. Иванов И.И. Основы программирования. — М.: Наука, 2020. — 256 с.", indent=0)
        _add_para(doc, "2. Петров П.П. Алгоритмы и структуры данных. — СПб.: Питер, 2019. — 312 с.", indent=0)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "BIB_NO_YEAR") == 0

    def test_entry_without_year(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
        _add_para(doc, "1. Иванов И.И. Основы программирования. — М.: Наука. — 256 с.", indent=0)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "BIB_NO_YEAR") >= 1

    def test_mixed_entries(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "СПИСОК ЛИТЕРАТУРЫ")
        _add_para(doc, "1. Хороший источник с годом 2021.", indent=0)
        _add_para(doc, "2. Плохой источник без указания года.", indent=0)
        _add_para(doc, "3. Ещё один нормальный, 2023 год.", indent=0)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "BIB_NO_YEAR") == 1

    def test_disabled_via_rules(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "СПИСОК ЛИТЕРАТУРЫ")
        _add_para(doc, "1. Источник без года.", indent=0)
        path = _save(doc, tmp_path)
        r = _check(path, {"checks": {"bibFormat": False}})
        assert _count(r, "BIB_NO_YEAR") == 0


# ═══════════════════════════════════════════
# Титульный лист и структура документа
# ═══════════════════════════════════════════

class TestTitlePageHandling:

    def test_title_page_not_checked_for_font_size(self, tmp_path):
        """Абзацы до первого заголовка (титульник) не проверяются на размер."""
        doc = _make_base()
        # Титульный лист с 12pt (корректно)
        _add_para(doc, "АСТРАХАНСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ", size=12,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=0)
        _add_para(doc, "Факультет ЦТиК", size=12,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=0)
        # Начало основного текста
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Основной текст работы размером 14 пунктов.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "FONT_SIZE_MISMATCH") == 0

    def test_body_after_heading_checked(self, tmp_path):
        """Основной текст после заголовка проверяется."""
        doc = _make_base()
        _add_para(doc, "Титульный лист 12pt", size=12, indent=0,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст с неправильным размером шрифта.", size=18)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "FONT_SIZE_MISMATCH") >= 1

    def test_centered_paragraph_no_indent_error(self, tmp_path):
        """Центрированные абзацы не проверяются на отступ."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Центрированный текст без отступа, это нормально.",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=0)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "INDENT_MISMATCH") == 0
        assert _count(r, "INDENT_NOT_SET") == 0


# ═══════════════════════════════════════════
# Подписи: секционная нумерация, тире
# ═══════════════════════════════════════════

class TestCaptionFormats:

    def test_section_numbering_figure_ok(self, tmp_path):
        """Рисунок 2.1 – Название — валидная подпись."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Результат на рисунке 2.1.")
        p = doc.add_paragraph()
        run = p.add_run("Рисунок 2.1 – Схема работы")
        run.font.size = Pt(14)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "FIGURE_CAPTION_FORMAT") == 0

    def test_hyphen_in_caption_is_error(self, tmp_path):
        """Дефис вместо тире в подписи — ошибка."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "На рисунке 1 показан результат.")
        p = doc.add_paragraph()
        run = p.add_run("Рисунок 1 - Результат")  # дефис!
        run.font.size = Pt(14)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "FIGURE_DASH_WRONG") >= 1

    def test_em_dash_in_caption_ok(self, tmp_path):
        """Длинное тире — валидно."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Результат на рисунке 1.")
        p = doc.add_paragraph()
        run = p.add_run("Рисунок 1 — Результат")
        run.font.size = Pt(14)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "FIGURE_DASH_WRONG") == 0
        assert _count(r, "FIGURE_CAPTION_FORMAT") == 0

    def test_en_dash_in_caption_ok(self, tmp_path):
        """Короткое тире – тоже валидно."""
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Результат на рисунке 1.")
        p = doc.add_paragraph()
        run = p.add_run("Рисунок 1 \u2013 Результат")  # en-dash
        run.font.size = Pt(14)
        path = _save(doc, tmp_path)
        r = _check(path)
        assert _count(r, "FIGURE_DASH_WRONG") == 0


# ═══════════════════════════════════════════
# Вердикт
# ═══════════════════════════════════════════

class TestVerdict:

    def test_good_verdict(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "СОДЕРЖАНИЕ")
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Корректно оформленный текст основной части работы.")
        _add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
        _add_para(doc, "1. Иванов И.И. Книга. — М.: Наука, 2020. — 100 с.", indent=0)
        path = _save(doc, tmp_path)
        r = _check(path)
        # Может быть несколько low, но не должно быть high
        assert r["summary"]["high"] == 0

    def test_bad_verdict_with_critical(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст с неправильным шрифтом Arial.", font="Arial")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert r["verdict"] == "bad"

    def test_ok_verdict_no_critical(self, tmp_path):
        doc = _make_base()
        _add_heading(doc, "СОДЕРЖАНИЕ")
        _add_heading(doc, "ВВЕДЕНИЕ")
        _add_para(doc, "Текст с правильным форматированием Times New Roman.")
        path = _save(doc, tmp_path)
        r = _check(path)
        assert r["verdict"] in ("good", "ok")


# ═══════════════════════════════════════════
# Edge cases для Python кода
# ═══════════════════════════════════════════

class TestCodeCheckerEdgeCases:

    def test_empty_file(self, tmp_path):
        path = _write_py(tmp_path, "empty.py", "")
        r = check_python_code(path, "empty.py")
        assert r["file_type"] == "python"
        assert r["total_issues"] == 0

    def test_only_comments(self, tmp_path):
        src = "# This is a comment\n# Another comment\n"
        path = _write_py(tmp_path, "comments.py", src)
        r = check_python_code(path, "comments.py")
        assert r["total_issues"] == 0

    def test_only_docstring(self, tmp_path):
        src = '"""Module docstring."""\n'
        path = _write_py(tmp_path, "docstring.py", src)
        r = check_python_code(path, "docstring.py")
        assert r["total_issues"] == 0

    def test_unicode_code(self, tmp_path):
        src = '"""Модуль с кириллицей."""\n\n\ndef привет():\n    """Функция приветствия."""\n    return "Привет мир"\n'
        path = _write_py(tmp_path, "unicode.py", src)
        r = check_python_code(path, "unicode.py")
        assert r["file_type"] == "python"
        # Кириллица не должна ломать чекер
        assert isinstance(r["total_issues"], int)

    def test_syntax_error_still_checks_style(self, tmp_path):
        src = "def foo(\n    x=1\n"  # незакрытая скобка
        path = _write_py(tmp_path, "broken.py", src)
        r = check_python_code(path, "broken.py")
        # flake8 может найти E999 или другие ошибки
        assert isinstance(r["total_issues"], int)
        assert "source_lines" in r

    def test_very_long_file(self, tmp_path):
        lines = ['"""Big module."""\n', "\n", "\n"]
        for i in range(200):
            lines.append(f"x_{i} = {i}\n")
        src = "".join(lines)
        path = _write_py(tmp_path, "big.py", src)
        r = check_python_code(path, "big.py")
        assert len(r["source_lines"]) == len(src.splitlines())

    def test_mixed_indentation(self, tmp_path):
        src = "def foo():\n\tx = 1\n    y = 2\n"
        path = _write_py(tmp_path, "mixed_indent.py", src)
        r = check_python_code(path, "mixed_indent.py")
        # Должна быть ошибка смешанной индентации
        assert r["total_issues"] > 0

    def test_trailing_whitespace(self, tmp_path):
        src = "x = 1   \ny = 2\n"
        path = _write_py(tmp_path, "trailing.py", src)
        r = check_python_code(path, "trailing.py")
        codes = {i["code"] for i in r["issues"]}
        assert "W291" in codes or "W293" in codes or r["total_issues"] > 0

    def test_multiple_blank_lines(self, tmp_path):
        src = '"""Module."""\n\n\n\n\ndef foo():\n    """Foo."""\n    pass\n'
        path = _write_py(tmp_path, "blanks.py", src)
        r = check_python_code(path, "blanks.py")
        codes = {i["code"] for i in r["issues"]}
        assert "E303" in codes  # слишком много пустых строк

    def test_no_newline_at_end(self, tmp_path):
        src = "x = 1"  # без \n в конце
        path = _write_py(tmp_path, "no_newline.py", src)
        r = check_python_code(path, "no_newline.py")
        codes = {i["code"] for i in r["issues"]}
        assert "W292" in codes  # no newline at end of file

    def test_verdict_good_for_clean(self, tmp_path):
        src = '"""Clean."""\n\n\ndef add(a, b):\n    """Add."""\n    return a + b\n'
        path = _write_py(tmp_path, "clean.py", src)
        r = check_python_code(path, "clean.py")
        assert r["verdict"] == "good"

    def test_verdict_bad_for_f401(self, tmp_path):
        src = "import os\nimport sys\n\nprint('hello')\n"
        path = _write_py(tmp_path, "unused.py", src)
        r = check_python_code(path, "unused.py")
        assert r["verdict"] == "bad"  # F401 = high severity

    def test_cyrillic_filename(self, tmp_path):
        src = '"""Тест."""\n'
        path = _write_py(tmp_path, "тест.py", src)
        r = check_python_code(path, "тест.py")
        assert r["filename"] == "тест.py"

    def test_single_line(self, tmp_path):
        src = "print('hello')\n"
        path = _write_py(tmp_path, "one.py", src)
        r = check_python_code(path, "one.py")
        assert len(r["source_lines"]) == 1
